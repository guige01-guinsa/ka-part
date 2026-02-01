from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from pathlib import Path
from datetime import datetime
from fastapi import Request

import json
from typing import Optional
from fastapi import Body

import io
import re
import tempfile

from app.db import db_conn
from app.auth import get_current_user
from app.converters import convert_with_soffice

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

FONT_PATH = os.path.join(os.getcwd(), "fonts", "NotoSansKR-Regular.ttf")
pdfmetrics.registerFont(TTFont("NotoKR", FONT_PATH))


router = APIRouter(prefix="/api/reports", tags=["reports"])


def _yyyymm_to_range(yyyymm: str):
    if not re.fullmatch(r"\d{4}-\d{2}", yyyymm or ""):
        raise HTTPException(status_code=400, detail="yyyymm must be YYYY-MM (e.g., 2026-01)")
    y, m = yyyymm.split("-")
    y = int(y)
    m = int(m)
    if m < 1 or m > 12:
        raise HTTPException(status_code=400, detail="Invalid month")

    # SQLite date range: [start, next_start)
    start = f"{y:04d}-{m:02d}-01"
    if m == 12:
        next_start = f"{y+1:04d}-01-01"
    else:
        next_start = f"{y:04d}-{m+1:02d}-01"
    return start, next_start


def _table_columns(table: str) -> set[str]:
    with db_conn() as db:
        cur = db.execute(f"PRAGMA table_info({table})")
        rows = cur.fetchall()
    return {r["name"] for r in rows}


def _monthly_work_data(yyyymm: str):
    start, next_start = _yyyymm_to_range(yyyymm)

    wo_cols = _table_columns("work_orders")
    ev_cols = _table_columns("events")

    # work_orders 기준일(월 집계 기준)
    # 1) created_at 있으면 created_at
    # 2) 없으면 events.CREATE created_at 사용
    has_wo_created = "created_at" in wo_cols
    has_wo_completed = "completed_at" in wo_cols
    has_wo_priority = "priority" in wo_cols  # 있으면 긴급 필터에 사용
    has_ev_created = "created_at" in ev_cols

    if not has_ev_created:
        raise HTTPException(status_code=500, detail="events.created_at missing (schema mismatch)")

    # 생성일/완료일 표현식(스키마에 따라 COALESCE 구성)
    created_expr = "wo.created_at" if has_wo_created else "ec.created_at"
    completed_expr = "wo.completed_at" if has_wo_completed else "ed.done_at"

    # 긴급 판정: priority 컬럼 있으면 우선, 없으면 title에 '긴급' 포함 여부로 임시 판정
    urgent_expr = "CASE WHEN wo.priority = 'URGENT' THEN 1 ELSE 0 END" if has_wo_priority else \
                  "CASE WHEN wo.title LIKE '%긴급%' THEN 1 ELSE 0 END"

    # location name 조인(있으면)
    has_locations = True
    try:
        _ = _table_columns("locations")
    except Exception:
        has_locations = False

    location_name_expr = "loc.name" if has_locations else "CAST(wo.location_id AS TEXT)"

    # 이벤트 기반 보조 서브쿼리 (created_at / done_at)
    # ec: CREATE 이벤트 시간
    # ed: DONE으로 바뀐 시각(STATUS_CHANGE to_status='DONE')
    sql = f"""
    WITH
    ec AS (
        SELECT entity_id AS work_id, MIN(created_at) AS created_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='CREATE'
        GROUP BY entity_id
    ),
    ed AS (
        SELECT entity_id AS work_id, MAX(created_at) AS done_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='STATUS_CHANGE' AND to_status='DONE'
        GROUP BY entity_id
    ),
    base AS (
        SELECT
            wo.id,
            wo.work_code,
            wo.status,
            wo.title,
            wo.category_id,
            wo.location_id,
            {created_expr} AS created_at,
            {completed_expr} AS completed_at,
            {urgent_expr} AS urgent_flag
        FROM work_orders wo
        LEFT JOIN ec ON ec.work_id = wo.id
        LEFT JOIN ed ON ed.work_id = wo.id
        WHERE {created_expr} >= ? AND {created_expr} < ?
    )
    SELECT
        (SELECT COUNT(*) FROM base) AS total_created,
        (SELECT COUNT(*) FROM base WHERE status='DONE') AS done_count,
        (SELECT COUNT(*) FROM base WHERE status!='DONE') AS not_done_count,
        (SELECT COUNT(*) FROM base WHERE urgent_flag=1) AS urgent_created,
        (SELECT COUNT(*) FROM base WHERE urgent_flag=1 AND status!='DONE') AS urgent_open,
        (SELECT COUNT(*) FROM base WHERE status='NEW') AS st_new,
        (SELECT COUNT(*) FROM base WHERE status='ASSIGNED') AS st_assigned,
        (SELECT COUNT(*) FROM base WHERE status='IN_PROGRESS') AS st_in_progress,
        (SELECT COUNT(*) FROM base WHERE status='REVIEW') AS st_review,
        (SELECT COUNT(*) FROM base WHERE status='APPROVED') AS st_approved,
        (SELECT COUNT(*) FROM base WHERE status='CANCELED') AS st_canceled
    ;
    """

    with db_conn() as db:
        cur = db.execute(sql, (start, next_start))
        head = cur.fetchone()

    # 위치별/상태별 집계
    if has_locations:
        loc_join = "LEFT JOIN locations loc ON loc.id = b.location_id"
    else:
        loc_join = ""

    sql_loc = f"""
    WITH
    ec AS (
        SELECT entity_id AS work_id, MIN(created_at) AS created_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='CREATE'
        GROUP BY entity_id
    ),
    ed AS (
        SELECT entity_id AS work_id, MAX(created_at) AS done_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='STATUS_CHANGE' AND to_status='DONE'
        GROUP BY entity_id
    ),
    b AS (
        SELECT
            wo.id,
            wo.status,
            wo.location_id,
            {"wo.created_at" if has_wo_created else "ec.created_at"} AS created_at,
            {urgent_expr} AS urgent_flag
        FROM work_orders wo
        LEFT JOIN ec ON ec.work_id = wo.id
        LEFT JOIN ed ON ed.work_id = wo.id
        WHERE ({"wo.created_at" if has_wo_created else "ec.created_at"}) >= ? AND ({"wo.created_at" if has_wo_created else "ec.created_at"}) < ?
    )
    SELECT
        {location_name_expr} AS location_name,
        COUNT(*) AS created,
        SUM(CASE WHEN b.status='DONE' THEN 1 ELSE 0 END) AS done,
        SUM(CASE WHEN b.status!='DONE' THEN 1 ELSE 0 END) AS open,
        SUM(CASE WHEN b.urgent_flag=1 THEN 1 ELSE 0 END) AS urgent
    FROM b
    {loc_join}
    GROUP BY location_name
    ORDER BY open DESC, urgent DESC, created DESC
    LIMIT 20;
    """

    with db_conn() as db:
        cur = db.execute(sql_loc, (start, next_start))
        loc_rows = cur.fetchall()

    # 대표 목록(미완료 상위 30건)
    sql_open = f"""
    WITH
    ec AS (
        SELECT entity_id AS work_id, MIN(created_at) AS created_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='CREATE'
        GROUP BY entity_id
    ),
    ed AS (
        SELECT entity_id AS work_id, MAX(created_at) AS done_at
        FROM events
        WHERE entity_type='WORK_ORDER' AND event_type='STATUS_CHANGE' AND to_status='DONE'
        GROUP BY entity_id
    )
    SELECT
        wo.id, wo.work_code, wo.status, wo.title,
        wo.category_id, wo.location_id,
        {"wo.created_at" if has_wo_created else "ec.created_at"} AS created_at,
        {urgent_expr} AS urgent_flag
    FROM work_orders wo
    LEFT JOIN ec ON ec.work_id = wo.id
    LEFT JOIN ed ON ed.work_id = wo.id
    WHERE ({"wo.created_at" if has_wo_created else "ec.created_at"}) >= ? AND ({"wo.created_at" if has_wo_created else "ec.created_at"}) < ?
      AND wo.status != 'DONE'
    ORDER BY urgent_flag DESC, created_at ASC
    LIMIT 30;
    """

    with db_conn() as db:
        cur = db.execute(sql_open, (start, next_start))
        open_rows = cur.fetchall()

    done_rate = 0.0
    if head["total_created"]:
        done_rate = round((head["done_count"] / head["total_created"]) * 100.0, 1)

    return {
        "yyyymm": yyyymm,
        "range": {"start": start, "next_start": next_start},
        "summary": {
            "total_created": head["total_created"],
            "done_count": head["done_count"],
            "not_done_count": head["not_done_count"],
            "done_rate_pct": done_rate,
            "urgent_created": head["urgent_created"],
            "urgent_open": head["urgent_open"],
        },
        "status_breakdown": {
            "NEW": head["st_new"],
            "ASSIGNED": head["st_assigned"],
            "IN_PROGRESS": head["st_in_progress"],
            "REVIEW": head["st_review"],
            "APPROVED": head["st_approved"],
            "CANCELED": head["st_canceled"],
            "DONE": head["done_count"],
        },
        "by_location": [dict(r) for r in loc_rows],
        "top_open": [dict(r) for r in open_rows],
    }


@router.get("/monthly-work")
async def monthly_work(request: Request, yyyymm: str):
    # 헤더 우선, 없으면 query(login=) 허용
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")

    # 사용자 유효성 검증(Unknown user 방지)
    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
    if not u:
        raise HTTPException(status_code=401, detail="Unknown user")
    if int(u["is_active"]) != 1:
        raise HTTPException(status_code=403, detail="Inactive user")

    return _monthly_work_data(yyyymm)

@router.get("/monthly-work.pdf")
async def monthly_work_pdf(request: Request, yyyymm: str):
    # 헤더 우선, 없으면 query(login=) 허용
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")

    # 사용자 유효성 검증(Unknown user 방지)
    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
    if not u:
        raise HTTPException(status_code=401, detail="Unknown user")
    if int(u["is_active"]) != 1:
        raise HTTPException(status_code=403, detail="Inactive user")

    data = _monthly_work_data(yyyymm)

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"reportlab missing: {e}")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    y = h - 20 * mm
    c.setFont("NotoKR", 14)
    c.drawString(20 * mm, y, f"Monthly Work Report  {data['yyyymm']}")
    y -= 10 * mm

    c.setFont("NotoKR", 10)
    s = data["summary"]
    c.drawString(20 * mm, y, f"Period: {data['range']['start']} ~ {data['range']['next_start']} (exclusive)")
    y -= 6 * mm
    c.drawString(20 * mm, y, f"Created: {s['total_created']}   Done: {s['done_count']}   Open: {s['not_done_count']}   DoneRate: {s['done_rate_pct']}%")
    y -= 6 * mm
    c.drawString(20 * mm, y, f"Urgent Created: {s['urgent_created']}   Urgent Open: {s['urgent_open']}")
    y -= 10 * mm

    c.setFont("NotoKR", 11)
    c.drawString(20 * mm, y, "Status Breakdown")
    y -= 6 * mm
    c.setFont("NotoKR", 10)
    sb = data["status_breakdown"]
    c.drawString(20 * mm, y, f"NEW {sb['NEW']} / ASSIGNED {sb['ASSIGNED']} / IN_PROGRESS {sb['IN_PROGRESS']} / REVIEW {sb['REVIEW']} / APPROVED {sb['APPROVED']} / DONE {sb['DONE']}")
    y -= 10 * mm

    c.setFont("NotoKR", 11)
    c.drawString(20 * mm, y, "Top Locations (by open)")
    y -= 6 * mm
    c.setFont("NotoKR", 9)
    for r in data["by_location"][:10]:
        line = f"{r.get('location_name','-')}: created {r.get('created',0)}, open {r.get('open',0)}, done {r.get('done',0)}, urgent {r.get('urgent',0)}"
        c.drawString(20 * mm, y, line[:120])
        y -= 5 * mm
        if y < 20 * mm:
            break

    y -= 6 * mm
    c.setFont("NotoKR", 11)
    c.drawString(20 * mm, y, "Top Open Works")
    y -= 6 * mm
    c.setFont("NotoKR", 9)
    for r in data["top_open"][:10]:
        code = r.get("work_code", f"#{r.get('id')}")
        st = r.get("status", "-")
        title = (r.get("title") or "")[:60]
        c.drawString(20 * mm, y, f"{code} [{st}] {title}")
        y -= 5 * mm
        if y < 20 * mm:
            break

    c.showPage()
    c.save()

    buf.seek(0)
    filename = f"monthly-work-{yyyymm}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'}
    )


def _build_docx(data: dict) -> io.BytesIO:
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx missing: {e}")

    doc = Document()
    doc.add_heading(f"Monthly Work Report {data['yyyymm']}", level=1)

    s = data["summary"]
    doc.add_paragraph(f"Period: {data['range']['start']} ~ {data['range']['next_start']} (exclusive)")
    doc.add_paragraph(
        f"Created: {s['total_created']} / Done: {s['done_count']} / Open: {s['not_done_count']} / DoneRate: {s['done_rate_pct']}%"
    )
    doc.add_paragraph(f"Urgent Created: {s['urgent_created']} / Urgent Open: {s['urgent_open']}")

    sb = data["status_breakdown"]
    doc.add_heading("Status Breakdown", level=2)
    doc.add_paragraph(
        f"NEW {sb['NEW']} / ASSIGNED {sb['ASSIGNED']} / IN_PROGRESS {sb['IN_PROGRESS']} / "
        f"REVIEW {sb['REVIEW']} / APPROVED {sb['APPROVED']} / DONE {sb['DONE']}"
    )

    doc.add_heading("Top Locations (by open)", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Location"
    hdr_cells[1].text = "Created"
    hdr_cells[2].text = "Open"
    hdr_cells[3].text = "Done"
    hdr_cells[4].text = "Urgent"
    for r in data["by_location"][:10]:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.get("location_name", "-"))
        row_cells[1].text = str(r.get("created", 0))
        row_cells[2].text = str(r.get("open", 0))
        row_cells[3].text = str(r.get("done", 0))
        row_cells[4].text = str(r.get("urgent", 0))

    doc.add_heading("Top Open Works", level=2)
    table2 = doc.add_table(rows=1, cols=3)
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Code"
    hdr2[1].text = "Status"
    hdr2[2].text = "Title"
    for r in data["top_open"][:10]:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(r.get("work_code", f"#{r.get('id')}"))
        row_cells[1].text = str(r.get("status", "-"))
        row_cells[2].text = str(r.get("title", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_xlsx(data: dict) -> io.BytesIO:
    try:
        from openpyxl import Workbook
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"openpyxl missing: {e}")

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"

    ws.append(["Monthly Work Report", data["yyyymm"]])
    ws.append(["Period", f"{data['range']['start']} ~ {data['range']['next_start']} (exclusive)"])
    s = data["summary"]
    ws.append(["Created", s["total_created"]])
    ws.append(["Done", s["done_count"]])
    ws.append(["Open", s["not_done_count"]])
    ws.append(["DoneRate%", s["done_rate_pct"]])
    ws.append(["Urgent Created", s["urgent_created"]])
    ws.append(["Urgent Open", s["urgent_open"]])

    sb = data["status_breakdown"]
    ws.append([])
    ws.append(["Status", "Count"])
    for k in ["NEW", "ASSIGNED", "IN_PROGRESS", "REVIEW", "APPROVED", "DONE"]:
        ws.append([k, sb.get(k, 0)])

    ws2 = wb.create_sheet(title="By Location")
    ws2.append(["Location", "Created", "Open", "Done", "Urgent"])
    for r in data["by_location"]:
        ws2.append([
            r.get("location_name", "-"),
            r.get("created", 0),
            r.get("open", 0),
            r.get("done", 0),
            r.get("urgent", 0),
        ])

    ws3 = wb.create_sheet(title="Top Open")
    ws3.append(["Code", "Status", "Title"])
    for r in data["top_open"]:
        ws3.append([
            r.get("work_code", f"#{r.get('id')}"),
            r.get("status", "-"),
            r.get("title", ""),
        ])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


@router.get("/monthly-work.docx")
async def monthly_work_docx(request: Request, yyyymm: str):
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")

    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
    if not u:
        raise HTTPException(status_code=401, detail="Unknown user")
    if int(u["is_active"]) != 1:
        raise HTTPException(status_code=403, detail="Inactive user")

    data = _monthly_work_data(yyyymm)
    buf = _build_docx(data)
    filename = f"monthly-work-{yyyymm}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/monthly-work.xlsx")
async def monthly_work_xlsx(request: Request, yyyymm: str):
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")

    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
    if not u:
        raise HTTPException(status_code=401, detail="Unknown user")
    if int(u["is_active"]) != 1:
        raise HTTPException(status_code=403, detail="Inactive user")

    data = _monthly_work_data(yyyymm)
    buf = _build_xlsx(data)
    filename = f"monthly-work-{yyyymm}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/monthly-work.hwp")
async def monthly_work_hwp(request: Request, yyyymm: str):
    """
    HWP 변환은 LibreOffice 또는 한컴 변환기 설치가 필요.
    """
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")

    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
    if not u:
        raise HTTPException(status_code=401, detail="Unknown user")
    if int(u["is_active"]) != 1:
        raise HTTPException(status_code=403, detail="Inactive user")

    data = _monthly_work_data(yyyymm)
    buf = _build_docx(data)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = f"{tmpdir}/monthly-work-{yyyymm}.docx"
        with open(docx_path, "wb") as f:
            f.write(buf.read())
        try:
            out_path = convert_with_soffice(docx_path, "hwp", tmpdir)
        except Exception as e:
            raise HTTPException(status_code=501, detail=f"HWP conversion failed: {e}")

        with open(out_path, "rb") as f:
            hwp_bytes = io.BytesIO(f.read())
            hwp_bytes.seek(0)

    filename = f"monthly-work-{yyyymm}.hwp"
    return StreamingResponse(
        hwp_bytes,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ============================================================
# Monthly Report Snapshot + Approval + PDF (1-page)
# ============================================================

def _require_login_from_header_or_query(request: Request) -> str:
    login = request.headers.get("X-User-Login") or request.query_params.get("login")
    if not login:
        raise HTTPException(status_code=401, detail="Missing X-User-Login header (or login query)")
    return login


async def _ensure_active_user(login: str) -> dict:
    """
    returns: {"login": str, "is_admin": bool, "roles": [str]}
    """
    with db_conn() as db:
        cur = db.execute("SELECT id, login, is_active FROM users WHERE login=?", (login,))
        u = cur.fetchone()
        if not u:
            raise HTTPException(status_code=401, detail="Unknown user")
        if int(u["is_active"]) != 1:
            raise HTTPException(status_code=403, detail="Inactive user")

        # roles via user_roles join
        cur = db.execute(
            """
            SELECT r.name AS role_name
            FROM user_roles ur
            JOIN roles r ON r.id = ur.role_id
            WHERE ur.user_id = ?
            """,
            (u["id"],),
        )
        rows = cur.fetchall()

    role_names = [r["role_name"] for r in rows] if rows else []

    # "role 기반" admin 판정 (현장 실무 안전장치로 login=admin도 admin 처리)
    # - roles 테이블에 '관리소장'이 존재하므로 기본적으로 관리소장을 admin으로 간주
    is_admin = (login == "admin") or any(
        (name or "").strip() in ("관리소장", "ADMIN", "admin") for name in role_names
    )

    return {"login": login, "roles": role_names, "is_admin": is_admin}

def _next_report_code(yyyymm: str) -> str:
    """
    MWR-YYYY-MM-0001 (월 단위 시퀀스)
    """
    prefix = f"MWR-{yyyymm}-"
    with db_conn() as db:
        cur = db.execute(
            "SELECT report_code FROM monthly_reports WHERE yyyymm=? AND report_code LIKE ? ORDER BY id DESC LIMIT 1",
            (yyyymm, prefix + "%"),
        )
        row = cur.fetchone()

    if not row:
        seq = 1
    else:
        last = row["report_code"][-4:]
        try:
            seq = int(last) + 1
        except:
            seq = 1

    return f"{prefix}{seq:04d}"


@router.post("/monthly-work/generate")
async def monthly_work_generate(request: Request, yyyymm: str):
    """
    월간 작업 실적 '스냅샷' 생성 + 보고서 번호 발급
    호출 예:
      /api/reports/monthly-work/generate?yyyymm=2026-01&login=admin
    """
    user = get_current_user(request)

    # 스냅샷 집계
    data = _monthly_work_data(yyyymm)

    # 보고서 코드 발급
    code = _next_report_code(yyyymm)

    with db_conn() as db:
        db.execute(
            """
            INSERT INTO monthly_reports (report_code, yyyymm, status, created_by, payload_json)
            VALUES (?, ?, 'DRAFT', ?, ?)
            """,
            (code, yyyymm, user["login"], json.dumps(data, ensure_ascii=False)),
        )
        db.commit()

        cur = db.execute("SELECT last_insert_rowid() AS id")
        rid = (cur.fetchone())["id"]

    return {"ok": True, "report_id": rid, "report_code": code, "status": "DRAFT"}



@router.get("/monthly-work/reports")
async def monthly_work_reports(request: Request, yyyymm: Optional[str] = None, limit: int = 50):
    """
    생성된 월간 보고서 목록
    """
    _ = get_current_user(request)



    q = """
    SELECT id, report_code, yyyymm, status, created_by, created_at, updated_at, submitted_at, approved_at, approved_by
    FROM monthly_reports
    """
    params = []
    if yyyymm:
        q += " WHERE yyyymm=?"
        params.append(yyyymm)
    q += " ORDER BY id DESC LIMIT ?"
    params.append(limit)

    with db_conn() as db:
        cur = db.execute(q, tuple(params))
        rows = cur.fetchall()

    return {"ok": True, "items": [dict(r) for r in rows]}


@router.get("/monthly-work/report")
async def monthly_work_report_get(request: Request, report_id: int):
    """
    월간 작업 실적 보고서 1건 조회 (스냅샷)
    """
    # 인증 확인 (역할 정보 포함)
    _ = get_current_user(request)

    with db_conn() as db:
        cur = db.execute(
            """
            SELECT
                id,
                report_code,
                yyyymm,
                status,
                created_by,
                created_at,
                updated_at,
                submitted_at,
                approved_at,
                approved_by,
                payload_json
            FROM monthly_reports
            WHERE id = ?
            """,
            (report_id,),
        )
        row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Report not found")

    data = dict(row)
    data["payload"] = json.loads(data.pop("payload_json"))

    return {"ok": True, "report": data}


@router.post("/monthly-work/report/status")
async def monthly_work_report_status(
    request: Request,
    report_id: int,
    to_status: str = Body(..., embed=True),
):
    """
    결재 상태 전이: DRAFT -> SUBMITTED -> APPROVED
    - APPROVED는 admin(역할)만 가능
    - APPROVED 이후는 잠금 유지
    """
    user = get_current_user(request)

    to_status = (to_status or "").upper().strip()
    if to_status not in ("SUBMITTED", "APPROVED"):
        raise HTTPException(status_code=400, detail="Invalid to_status")

    with db_conn() as db:
        cur = db.execute("SELECT id, status FROM monthly_reports WHERE id=?", (report_id,))
        r = cur.fetchone()
        if not r:
            raise HTTPException(status_code=404, detail="Report not found")

        cur_status = r["status"]

        allowed = {
            "DRAFT": {"SUBMITTED"},
            "SUBMITTED": {"APPROVED"},
            "APPROVED": set(),
        }

        if to_status not in allowed.get(cur_status, set()):
            raise HTTPException(status_code=400, detail=f"Invalid transition {cur_status} -> {to_status}")

        # 🔒 승인(APPROVED)은 admin(역할)만
        if to_status == "APPROVED" and not user["is_admin"]:
            raise HTTPException(status_code=403, detail="승인(APPROVED)은 관리자만 가능합니다.")

        if to_status == "SUBMITTED":
            db.execute(
                "UPDATE monthly_reports SET status='SUBMITTED', submitted_at=datetime('now') WHERE id=?",
                (report_id,),
            )
        elif to_status == "APPROVED":
            db.execute(
                "UPDATE monthly_reports SET status='APPROVED', approved_at=datetime('now'), approved_by=? WHERE id=?",
                (user["login"], report_id),
            )

        db.commit()

    return {"ok": True, "from": cur_status, "to": to_status, "is_admin": user["is_admin"], "roles": user["roles"]}


@router.get("/monthly-work/report.pdf")
async def monthly_work_report_pdf(request: Request, report_id: int):
    """
    스냅샷 기준 1페이지 PDF (결재란 포함)
    호출 예:
      /api/reports/monthly-work/report.pdf?report_id=1&login=admin
    """
    _ = get_current_user(request)


    with db_conn() as db:
        cur = db.execute(
            "SELECT report_code, yyyymm, status, created_by, created_at, submitted_at, approved_at, approved_by, payload_json "
            "FROM monthly_reports WHERE id=?",
            (report_id,),
        )
        row = cur.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Report not found")

    meta = dict(row)
    data = json.loads(meta["payload_json"])

    # reportlab
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"reportlab missing: {e}")

    # 한글 폰트 (프로젝트 루트/fonts)
    import os
    font_path_ttf = os.path.join(os.getcwd(), "fonts", "NotoSansKR-Regular.ttf")
    if not os.path.exists(font_path_ttf):
        raise HTTPException(status_code=500, detail=f"Font not found: {font_path_ttf}")
    pdfmetrics.registerFont(TTFont("NotoKR", font_path_ttf))

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4

    # ---- Header
    y = h - 18 * mm
    c.setFont("NotoKR", 14)
    c.drawString(18 * mm, y, f"월간 작업 실적 (입대의 제출용 / 내부)")
    y -= 7 * mm

    c.setFont("NotoKR", 10)
    c.drawString(18 * mm, y, f"보고서번호: {meta['report_code']}    월: {meta['yyyymm']}    상태: {meta['status']}")
    y -= 5 * mm
    c.drawString(18 * mm, y, f"작성: {meta['created_by']}  ({meta['created_at']})")
    y -= 5 * mm
    c.drawString(18 * mm, y, f"제출: {meta.get('submitted_at') or '-'}    승인: {meta.get('approved_by') or '-'}  ({meta.get('approved_at') or '-'})")
    y -= 8 * mm

    # ---- Summary
    s = data["summary"]
    rng = data["range"]

    c.setFont("NotoKR", 11)
    c.drawString(18 * mm, y, "요약")
    y -= 6 * mm
    c.setFont("NotoKR", 10)
    c.drawString(18 * mm, y, f"기간: {rng['start']} ~ {rng['next_start']} (미포함)")
    y -= 5 * mm
    c.drawString(18 * mm, y, f"생성: {s['total_created']} / 완료(DONE): {s['done_count']} / 미완료: {s['not_done_count']} / 완료율: {s['done_rate_pct']}%")
    y -= 5 * mm
    c.drawString(18 * mm, y, f"긴급 생성: {s['urgent_created']} / 긴급 미완료: {s['urgent_open']}")
    y -= 8 * mm

    # ---- Status breakdown
    sb = data["status_breakdown"]
    c.setFont("NotoKR", 11)
    c.drawString(18 * mm, y, "상태 분포")
    y -= 6 * mm
    c.setFont("NotoKR", 10)
    c.drawString(
        18 * mm, y,
        f"NEW {sb.get('NEW',0)} / ASSIGNED {sb.get('ASSIGNED',0)} / IN_PROGRESS {sb.get('IN_PROGRESS',0)} / REVIEW {sb.get('REVIEW',0)} / APPROVED {sb.get('APPROVED',0)} / DONE {sb.get('DONE',0)}"
    )
    y -= 10 * mm

    # ---- Tables
    c.setFont("NotoKR", 11)
    c.drawString(18 * mm, y, "위치별(미완료 우선) 상위 10")
    y -= 6 * mm
    c.setFont("NotoKR", 9)
    c.drawString(18 * mm, y, "위치 / 생성 / 미완료 / 완료 / 긴급")
    y -= 5 * mm

    for r in data.get("by_location", [])[:10]:
        line = f"{r.get('location_name','-')} / {r.get('created',0)} / {r.get('open',0)} / {r.get('done',0)} / {r.get('urgent',0)}"
        c.drawString(18 * mm, y, line[:110])
        y -= 5 * mm
        if y < 55 * mm:
            break

    y -= 6 * mm
    c.setFont("NotoKR", 11)
    c.drawString(18 * mm, y, "미완료 TOP 10")
    y -= 6 * mm
    c.setFont("NotoKR", 9)
    top_open = data.get("top_open", [])[:10]
    if not top_open:
        c.drawString(18 * mm, y, "미완료 없음")
        y -= 5 * mm
    else:
        for r in top_open:
            code = r.get("work_code", f"#{r.get('id','')}")
            st = r.get("status", "-")
            title = (r.get("title") or "")[:60]
            c.drawString(18 * mm, y, f"{code} [{st}] {title}")
            y -= 5 * mm
            if y < 55 * mm:
                break

    # ---- Approval box (bottom)
    y = 40 * mm
    c.setFont("NotoKR", 10)
    c.drawString(18 * mm, y + 18, "결재")
    c.rect(18 * mm, y, 170 * mm, 16 * mm, stroke=1, fill=0)
    c.drawString(22 * mm, y + 5, "작성")
    c.drawString(70 * mm, y + 5, "검토")
    c.drawString(120 * mm, y + 5, "승인")
    c.line(55 * mm, y, 55 * mm, y + 16 * mm)
    c.line(105 * mm, y, 105 * mm, y + 16 * mm)

    c.showPage()
    c.save()

    buf.seek(0)
    filename = f"{meta['report_code']}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{filename}"'}
    )
