# app.py — Android(Pydroid3) 호환 All-in-One 아파트 관리 클라이언트
# 기능: 일검침 일지(전력/급수/열량/유량) + 민원/고장 + 업무파일 + 설정
# 보조: 누락 컬럼 자동 마이그레이션, CSV 내보내기, 월별 집계, 일괄 재계산, 음성 입력
# 주의: 같은 폰의 크롬에서 http://127.0.0.1:8000/ 로 접속

# 1️⃣ import 구문 (파일 상단)
from blueprints.tool_search.main import tool_search_bp
import os, math, json, mimetypes, uuid
import requests  # ← 카카오 REST 호출
from datetime import datetime, date, time
from typing import Optional, List

from flask import (
    Flask, request, jsonify, redirect, url_for,
    render_template_string, flash, send_file
)
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func, text

from sqlalchemy.exc import OperationalError

SUBTABS = """
<div class="mb-3">
  <ul class="nav nav-pills">
    <li class="nav-item">
      <a class="nav-link {{ 'active' if request.path.startswith('/files') else '' }}"
         href="{{ url_for('ui_files') }}">📂 파일보관</a>
    </li>
    <li class="nav-item">
      <a class="nav-link {{ 'active' if request.path.startswith('/compare') else '' }}"
         href="{{ url_for('ui_compare') }}">⚡ 비교견적</a>
    </li>
  </ul>
</div>
"""

def ensure_compare_schema():
    """비교견적 관련 테이블이 없으면 즉시 생성/보강"""
    with app.app_context():
        db.create_all()
        try:
            # 존재 확인 겸 간단 조회
            db.session.execute(text("SELECT 1 FROM compare_set LIMIT 1"))
        except Exception:
            # 초창기 DB에는 없을 수 있음 → 보강 루틴
            auto_migrate_columns()

# ───────────────────────────────────────────────────────────────────
# 1) Flask & SQLite 초기 설정 (instance 폴더를 DB/업로드 저장소로 사용)
#    - Pydroid3에서 쓰기 권한 보장
# ───────────────────────────────────────────────────────────────────
app = Flask(__name__, instance_relative_config=True)
os.makedirs(app.instance_path, exist_ok=True)
UPLOAD_DIR = os.path.join(app.instance_path, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

db_path = os.path.join(app.instance_path, "apartment.db")

# ===============================
# Database URL normalize
# ===============================

def _normalize_database_url(url: str) -> str:
    # postgres:// -> postgresql:// 보정
    if url.startswith("postgres://"):
        url = "postgresql://" + url[len("postgres://"):]
    # Render Postgres SSL 요구 대응
    if url.startswith("postgresql://") and "sslmode=" not in url:
        join = "&" if "?" in url else "?"
        url = url + join + "sslmode=require"
    return url


# ===============================
# Database configuration
# ===============================

# 기본 SQLite (DATABASE_URL 없을 때)
db_path = os.path.join(app.instance_path, "apartment.db")
default_sqlite = "sqlite:///" + db_path

raw_db_url = os.environ.get("DATABASE_URL", "").strip()

if raw_db_url:
    db_url = _normalize_database_url(raw_db_url)
else:
    db_url = default_sqlite

app.config["SQLALCHEMY_DATABASE_URI"] = db_url

# 엔진 옵션 분기 (SQLite만 check_same_thread 필요)
if app.config["SQLALCHEMY_DATABASE_URI"].startswith("sqlite:///"):
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
        "connect_args": {"check_same_thread": False}
    }
else:
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {}



else:
    # Render 우선: DATABASE_URL -> 없으면 SQLite(instance)
db_path = os.path.join(app.instance_path, "apartment.db")
default_sqlite = "sqlite:///" + db_path

db_url = os.environ.get("DATABASE_URL", "").strip()
if db_url.startswith("postgres://"):
    db_url = db_url.replace("postgres://", "postgresql://", 1)

app.config["SQLALCHEMY_DATABASE_URI"] = db_url or default_sqlite

# Postgres에서는 check_same_thread 옵션 넣으면 안 됨 → SQLite일 때만
if app.config["SQLALCHEMY_DATABASE_URI"].startswith("sqlite:///"):
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"connect_args": {"check_same_thread": False}}
else:
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {}


app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# SQLite에서만 check_same_thread 옵션이 의미 있음
if app.config["SQLALCHEMY_DATABASE_URI"].startswith("sqlite:///"):
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"connect_args": {"check_same_thread": False}}
else:
    # Postgres에서는 pool_pre_ping 정도만(선택)
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"pool_pre_ping": True}

app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
# 안드로이드 환경에서 같은 스레드 체크로 생기는 경고 방지
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"connect_args": {"check_same_thread": False}}
app.config["SECRET_KEY"] = "replace_me_for_forms"

db = SQLAlchemy(app)

app.register_blueprint(tool_search_bp, url_prefix="/ka-part/ts")

# ───────────────────────────────────────────────────────────────────
# 2) 유틸 함수들 (파서/반올림/전력 계산)
# ───────────────────────────────────────────────────────────────────
def parse_date(s):
    """'YYYY-MM-DD' → date, 빈 값은 None"""
    if not s: return None
    if isinstance(s, date): return s
    return datetime.strptime(s, "%Y-%m-%d").date()

def parse_time(s):
    """'HH:MM' → time, 빈 값은 None"""
    if not s: return None
    if isinstance(s, time): return s
    return datetime.strptime(s, "%H:%M").time()

def parse_float(s):
    """문자열 → float, 빈 값/잘못된 값은 None"""
    if s in (None, "", "null"): return None
    try: return float(str(s).replace(",", ""))
    except Exception: return None

def r3(x):
    """소수점 3자리 반올림"""
    if x is None: return None
    try: return round(float(x), 3)
    except Exception: return None

def calc_kw(voltage, current, pf, *, is_kv=False):
    """
    3상 유효전력 kW = √3 * V * I * PF / 1000
    - 고압측: 전압이 kV 단위 → is_kv=True 로 받아 V로 변환 후 계산
    """
    if voltage is None or current is None or pf in (None, "", "null"):
        return None
    try:
        v = float(voltage) * (1000 if is_kv else 1)
        i = float(current)
        pf = float(pf)
        kw = math.sqrt(3) * v * i * pf / 1000.0
        return r3(kw)
    except Exception:
        return None

# 기본 점검자(드롭다운)
OPERATORS = ["이상석", "이창희", "신충기"]
DEFAULT_OPERATOR = OPERATORS[0]

# ───────────────────────────────────────────────────────────────────
# 3) 모델 정의
#    - Settings: 사용량 보정 계수/요금/카카오 전송 설정
#    - SubstationLog: 일검침/설비 일지 (누적/일사용량 포함)
#    - WorkFile: 업무파일 저장소
#    - Complaint: 민원/고장 접수
# ───────────────────────────────────────────────────────────────────
class Settings(db.Model):
    __tablename__ = "settings"
    id = db.Column(db.Integer, primary_key=True)
    public_base_url = db.Column(db.String)  # 업로드/첨부의 공개 URL prefix(선택)

    # 사용량 보정 계수(일 사용량 계산시 곱함)
    hv_factor = db.Column(db.Float)       # 고압 사용량(요청: 1800 배)
    ind_factor = db.Column(db.Float)      # 산업용 사용량(요청: 30 배)
    street_factor = db.Column(db.Float)   # 가로등 사용량(기본 1 배)

    # 급수/열량/유량 보정(선택)
    water_factor = db.Column(db.Float, default=1.0)
    heat_factor  = db.Column(db.Float, default=1.0)
    flow_factor  = db.Column(db.Float, default=1.0)

    # 고지/배분(옵션)
    tariff_per_kwh = db.Column(db.Float, default=0.0)
    base_charge    = db.Column(db.Float, default=0.0)
    allocation_method = db.Column(db.String, default="equal")

    # 카카오 전송 설정(옵션)
    kakao_rest_key = db.Column(db.String)
    kakao_access_token = db.Column(db.String)
    kakao_friend_uuid  = db.Column(db.String)

    @staticmethod
    def get():
        """ID=1의 설정 레코드를 항상 보장(없으면 생성)"""
        row = Settings.query.get(1)
        if not row:
            row = Settings(
                id=1,
                hv_factor=1800.0, ind_factor=30.0, street_factor=1.0,
                water_factor=1.0, heat_factor=1.0, flow_factor=1.0,
                tariff_per_kwh=0.0, base_charge=0.0, allocation_method="equal",
                public_base_url=None,
            )
            db.session.add(row); db.session.commit()
        return row

class SubstationLog(db.Model):
    """
    일검침/설비일지
    - 전력: 고압 수전, 저압 3회선, 누적/일사용량
    - 설비: 급수/열량/유량 누적/일사용량, 각종 온도
    """
    __tablename__ = "substation_log"
    id = db.Column(db.Integer, primary_key=True)

    # 공통
    log_date = db.Column(db.Date, nullable=False)   # 일지 날짜
    log_time = db.Column(db.Time)                   # 일지 시각
    operator = db.Column(db.String(32), default=DEFAULT_OPERATOR)  # 점검자

    # 고압 수전측 (전압 kV 입력 → 자동 kW 계산)
    incomer_voltage = db.Column(db.Float)    # kV
    incomer_curr    = db.Column(db.Float)    # A
    vcb_p_factor    = db.Column(db.Float)    # 역률
    electric_energy = db.Column(db.Float)    # kW(자동계산)

    # 저압측 (V 입력 → 자동 kW 계산)
    lv1_v = db.Column(db.Float); lv1_a = db.Column(db.Float); lv1_kw = db.Column(db.Float)
    lv2_v = db.Column(db.Float); lv2_a = db.Column(db.Float); lv2_kw = db.Column(db.Float)
    lv3_v = db.Column(db.Float); lv3_a = db.Column(db.Float); lv3_kw = db.Column(db.Float)
    power_factor = db.Column(db.Float)       # 저압 공통 역률

    # 전력 누적/일사용량
    hv_acc_kwh   = db.Column(db.Float)   # 누적 고압 유효전력
    ind_acc_kwh  = db.Column(db.Float)   # 누적 산업용 유효전력
    str_acc_kwh  = db.Column(db.Float)   # 누적 가로등 유효전력
    hv_use_kwh   = db.Column(db.Float)   # 일 사용량(보정 반영) ← (오늘-전일)*계수
    ind_use_kwh  = db.Column(db.Float)
    str_use_kwh  = db.Column(db.Float)

    # 급수/열량/유량 (누적/일사용)
    acc_water = db.Column(db.Float); day_water = db.Column(db.Float)
    acc_heat  = db.Column(db.Float); day_heat  = db.Column(db.Float)
    acc_flow  = db.Column(db.Float); day_flow  = db.Column(db.Float)

    # 온도(설비)
    hst = db.Column(db.Float); hrt = db.Column(db.Float)
    lst = db.Column(db.Float); lrt = db.Column(db.Float)
    TR1 = db.Column(db.Float); TR2 = db.Column(db.Float)

    # 기타
    TR3 = db.Column(db.Float); winding_temp = db.Column(db.Float)
    event = db.Column(db.String(160), default="")   # 특이사항
    remarks = db.Column(db.Text, default="")        # 비고
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    # API/CSV/폼에 쓰기 좋게 dict 변환(반올림)
    def to_dict(self):
        f = lambda x: None if x is None else r3(x)
        return {
            "id": self.id,
            "log_date": self.log_date.isoformat(),
            "log_time": self.log_time.strftime("%H:%M") if self.log_time else None,
            "operator": self.operator,
            "incomer_voltage": f(self.incomer_voltage), "incomer_curr": f(self.incomer_curr), "vcb_p_factor": f(self.vcb_p_factor),
            "electric_energy": f(self.electric_energy),
            "lv1_v": f(self.lv1_v), "lv1_a": f(self.lv1_a), "lv1_kw": f(self.lv1_kw),
            "lv2_v": f(self.lv2_v), "lv2_a": f(self.lv2_a), "lv2_kw": f(self.lv2_kw),
            "lv3_v": f(self.lv3_v), "lv3_a": f(self.lv3_a), "lv3_kw": f(self.lv3_kw),
            "power_factor": f(self.power_factor),
            "hv_acc_kwh": f(self.hv_acc_kwh), "hv_use_kwh": f(self.hv_use_kwh),
            "ind_acc_kwh": f(self.ind_acc_kwh), "ind_use_kwh": f(self.ind_use_kwh),
            "str_acc_kwh": f(self.str_acc_kwh), "str_use_kwh": f(self.str_use_kwh),
            "acc_water": f(self.acc_water), "day_water": f(self.day_water),
            "acc_heat":  f(self.acc_heat),  "day_heat":  f(self.day_heat),
            "acc_flow":  f(self.acc_flow),  "day_flow":  f(self.day_flow),
            "hst": f(self.hst), "hrt": f(self.hrt), "lst": f(self.lst), "lrt": f(self.lrt),
            "TR1": f(self.TR1), "TR2": f(self.TR2),
            "TR3": f(self.TR3), "winding_temp": f(self.winding_temp),
            "event": self.event, "remarks": self.remarks,
        }

class WorkFile(db.Model):
    """업무 파일 보관/분류/전송(메타 정보)"""
    __tablename__ = "work_file"
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    title = db.Column(db.String(140)); description = db.Column(db.Text)
    category = db.Column(db.String(40)); tags = db.Column(db.String(140))
    filename = db.Column(db.String(300)); ext = db.Column(db.String(15))
    size = db.Column(db.Integer); uploader = db.Column(db.String(40))

class Complaint(db.Model):
    """민원/고장 접수 (간이 자동 분류 + 미디어 첨부)"""
    __tablename__ = "complaint"
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    name = db.Column(db.String(40)); unit = db.Column(db.String(40))
    phone = db.Column(db.String(40)); channel = db.Column(db.String(20))
    text = db.Column(db.Text)
    category = db.Column(db.String(40)); confidence = db.Column(db.Float)
    tags = db.Column(db.String(140))
    status = db.Column(db.String(20), default="접수")
    priority = db.Column(db.String(20), default="보통")
    assigned_to = db.Column(db.String(20), default="전기과장")
    media_filename = db.Column(db.String(300)); media_type = db.Column(db.String(20))
# ───────────────────────────────────────────────────────────────────
# 3-추가) 비교견적 모델
#  - CompareSet: 비교 단위(제목/의뢰자/메모)
#  - Vendor: 업체 기본정보
#  - Item: 품명/규격/수량/단위
#  - VendorPrice: 업체별-품목별 단가(공급가 기준, 부가세는 집계시 계산)
# ───────────────────────────────────────────────────────────────────
class CompareSet(db.Model):
    __tablename__ = "compare_set"
    id = db.Column(db.Integer, primary_key=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    title = db.Column(db.String(140), nullable=False)
    requester = db.Column(db.String(80))
    memo = db.Column(db.Text)
    vat_rate = db.Column(db.Float, default=0.1)  # 10% 기본

class Vendor(db.Model):
    __tablename__ = "vendor"
    id = db.Column(db.Integer, primary_key=True)
    compare_id = db.Column(db.Integer, db.ForeignKey("compare_set.id"), nullable=False, index=True)
    name = db.Column(db.String(120), nullable=False)
    contact = db.Column(db.String(120))
    phone = db.Column(db.String(40))
    note = db.Column(db.String(200))

class Item(db.Model):
    __tablename__ = "cmp_item"
    id = db.Column(db.Integer, primary_key=True)
    compare_id = db.Column(db.Integer, db.ForeignKey("compare_set.id"), nullable=False, index=True)
    name = db.Column(db.String(160), nullable=False)
    spec = db.Column(db.String(200))
    unit = db.Column(db.String(20), default="EA")
    qty = db.Column(db.Float, default=1.0)

class VendorPrice(db.Model):
    __tablename__ = "vendor_price"
    id = db.Column(db.Integer, primary_key=True)
    vendor_id = db.Column(db.Integer, db.ForeignKey("vendor.id"), nullable=False, index=True)
    item_id = db.Column(db.Integer, db.ForeignKey("cmp_item.id"), nullable=False, index=True)
    unit_price = db.Column(db.Float, nullable=True)  # 공급가(단가). None이면 미제시
    __table_args__ = (
        db.UniqueConstraint("vendor_id","item_id", name="uq_vendor_item"),
    )
# ───────────────────────────────────────────────────────────────────
# 4) DB 자동 마이그레이션 (기존 DB에 누락 컬럼이 있어도 안전하게 보강)
# ───────────────────────────────────────────────────────────────────
def table_exists(conn, name):
    names = {r[0] for r in conn.execute(text("SELECT name FROM sqlite_master WHERE type='table'"))}
    return name in names

def get_cols(conn, table):
    return {row[1] for row in conn.execute(text(f"PRAGMA table_info({table})"))}

def ensure_column(conn, table, name, ddl):
    """테이블에 컬럼이 없으면 ALTER TABLE ADD COLUMN 실행"""
    if not table_exists(conn, table): return
    if name not in get_cols(conn, table):
        conn.execute(text(f"ALTER TABLE {table} ADD COLUMN {name} {ddl}"))

def auto_migrate_columns():
    """필요 테이블 생성 + 모든 사용 칼럼 보강 (SQLite 전용)"""
    if db.engine.dialect.name != "sqlite":
        return
    # 이하 기존 코드 그대로...

    with db.engine.begin() as conn:
        # 최소 테이블 생성
        conn.execute(text("CREATE TABLE IF NOT EXISTS settings (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS substation_log (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS work_file (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS complaint (id INTEGER PRIMARY KEY)"))

        # settings 보강
        for name, ddl in [
            ("public_base_url","TEXT"),
            ("hv_factor","REAL"), ("ind_factor","REAL"), ("street_factor","REAL"),
            ("water_factor","REAL"), ("heat_factor","REAL"), ("flow_factor","REAL"),
            ("tariff_per_kwh","REAL"), ("base_charge","REAL"),
            ("allocation_method","TEXT"),
            ("kakao_rest_key","TEXT"), ("kakao_access_token","TEXT"), ("kakao_friend_uuid","TEXT"),
        ]:
            ensure_column(conn, "settings", name, ddl)

        # substation_log 보강 (폼/API와 1:1 매핑되는 모든 필드)
        for name, ddl in [
            ("log_date","DATE"), ("log_time","TIME"), ("operator","TEXT"),
            ("incomer_voltage","REAL"), ("incomer_curr","REAL"), ("vcb_p_factor","REAL"),
            ("electric_energy","REAL"),
            ("lv1_v","REAL"), ("lv1_a","REAL"), ("lv1_kw","REAL"),
            ("lv2_v","REAL"), ("lv2_a","REAL"), ("lv2_kw","REAL"),
            ("lv3_v","REAL"), ("lv3_a","REAL"), ("lv3_kw","REAL"),
            ("power_factor","REAL"),
            ("hv_acc_kwh","REAL"), ("ind_acc_kwh","REAL"), ("str_acc_kwh","REAL"),
            ("hv_use_kwh","REAL"), ("ind_use_kwh","REAL"), ("str_use_kwh","REAL"),
            ("acc_water","REAL"), ("day_water","REAL"),
            ("acc_heat","REAL"),  ("day_heat","REAL"),
            ("acc_flow","REAL"),  ("day_flow","REAL"),
            ("hst","REAL"), ("hrt","REAL"), ("lst","REAL"), ("lrt","REAL"),
            ("TR1","REAL"), ("TR2","REAL"),
            ("TR3","REAL"), ("winding_temp","REAL"),
            ("event","TEXT"), ("remarks","TEXT"), ("created_at","TEXT"),
        ]:
            ensure_column(conn, "substation_log", name, ddl)

        # work_file 보강
        for name, ddl in [
            ("created_at","TEXT"), ("title","TEXT"), ("description","TEXT"),
            ("category","TEXT"), ("tags","TEXT"),
            ("filename","TEXT"), ("ext","TEXT"), ("size","INTEGER"),
            ("uploader","TEXT"),
        ]:
            ensure_column(conn, "work_file", name, ddl)

        # complaint 보강
        for name, ddl in [
            ("created_at","TEXT"), ("name","TEXT"), ("unit","TEXT"),
            ("phone","TEXT"), ("channel","TEXT"), ("text","TEXT"),
            ("category","TEXT"), ("confidence","REAL"), ("tags","TEXT"),
            ("status","TEXT"), ("priority","TEXT"), ("assigned_to","TEXT"),
            ("media_filename","TEXT"), ("media_type","TEXT"),
        ]:
            ensure_column(conn, "complaint", name, ddl)

# compare_set / vendor / cmp_item / vendor_price 테이블 최소 생성
        conn.execute(text("CREATE TABLE IF NOT EXISTS compare_set (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS vendor (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS cmp_item (id INTEGER PRIMARY KEY)"))
        conn.execute(text("CREATE TABLE IF NOT EXISTS vendor_price (id INTEGER PRIMARY KEY)"))

        # compare_set 보강
        for name, ddl in [
            ("created_at","TEXT"), ("title","TEXT"), ("requester","TEXT"),
            ("memo","TEXT"), ("vat_rate","REAL"),
        ]:
            ensure_column(conn, "compare_set", name, ddl)

        # vendor 보강
        for name, ddl in [
            ("compare_id","INTEGER"), ("name","TEXT"), ("contact","TEXT"),
            ("phone","TEXT"), ("note","TEXT"),
        ]:
            ensure_column(conn, "vendor", name, ddl)

        # cmp_item 보강
        for name, ddl in [
            ("compare_id","INTEGER"), ("name","TEXT"), ("spec","TEXT"),
            ("unit","TEXT"), ("qty","REAL"),
        ]:
            ensure_column(conn, "cmp_item", name, ddl)

        # vendor_price 보강
        for name, ddl in [
            ("vendor_id","INTEGER"), ("item_id","INTEGER"), ("unit_price","REAL"),
        ]:
            ensure_column(conn, "vendor_price", name, ddl)
# 5) 자동 계산 로직
#    - 수전/저압 kW 자동계산
#    - 일사용량 = (금일 누적 - 전일 누적) * 계수(설정값)
#    - 음수 방지, 소수점 3자리 반올림
# ───────────────────────────────────────────────────────────────────
def compute_auto_fields(row: SubstationLog, prev: Optional[SubstationLog], s: Settings):
    # 5-1) 수전 및 저압 kW 계산
    row.electric_energy = calc_kw(row.incomer_voltage, row.incomer_curr, row.vcb_p_factor, is_kv=True)
    row.lv1_kw = calc_kw(row.lv1_v, row.lv1_a, row.power_factor, is_kv=False)
    row.lv2_kw = calc_kw(row.lv2_v, row.lv2_a, row.power_factor, is_kv=False)
    row.lv3_kw = calc_kw(row.lv3_v, row.lv3_a, row.power_factor, is_kv=False)

    # 5-2) 일 사용량 계산 도우미
    def diff_mul(today, yest, factor):
        if today is None or yest is None: return None
        val = (float(today) - float(yest)) * float(factor)
        return r3(max(val, 0.0))  # 누계 감소(계기 리셋 등)는 0으로 처리

    # 5-3) 전일 데이터가 있으면 일 사용량 산출
    if prev:
        row.hv_use_kwh  = diff_mul(row.hv_acc_kwh,  prev.hv_acc_kwh,  s.hv_factor or 1.0)
        row.ind_use_kwh = diff_mul(row.ind_acc_kwh, prev.ind_acc_kwh, s.ind_factor or 1.0)
        row.str_use_kwh = diff_mul(row.str_acc_kwh, prev.str_acc_kwh, s.street_factor or 1.0)

        row.day_water = diff_mul(row.acc_water, prev.acc_water, s.water_factor or 1.0)
        row.day_heat  = diff_mul(row.acc_heat,  prev.acc_heat,  s.heat_factor  or 1.0)
        row.day_flow  = diff_mul(row.acc_flow,  prev.acc_flow,  s.flow_factor  or 1.0)
    else:
        row.hv_use_kwh = row.ind_use_kwh = row.str_use_kwh = None
        row.day_water = row.day_heat = row.day_flow = None

# ───────────────────────────────────────────────────────────────────
# 6) 공통 라우트
# ───────────────────────────────────────────────────────────────────
@app.route("/health")
def health(): return jsonify(status="ok")

@app.route("/")
def home():
    return redirect(url_for("ui_apps"))




# ───────────────────────────────────────────────────────────────────
# 7) UI 기본 레이아웃 (부트스트랩 + 탭 네비)
#    - 음성 입력 지원 버튼 포함
# ───────────────────────────────────────────────────────────────────
BASE = """
<!doctype html><html lang="ko"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{{ title or '아파트 관리' }}</title>
<link rel="manifest" href="/ka-part/manifest.webmanifest">
<meta name="theme-color" content="#111111">

<link rel="icon" sizes="192x192" href="{{ url_for('static', filename='icons/icon-192.png') }}">
<link rel="icon" sizes="512x512" href="{{ url_for('static', filename='icons/icon-512.png') }}">
<link rel="apple-touch-icon" href="{{ url_for('static', filename='icons/icon-192.png') }}">

<link rel="icon" href="{{ url_for('static', filename='favicon.ico') }}">

<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<style>body{background:#f7f8fb;padding-bottom:70px}.table-sm td,.table-sm th{padding:.35rem .5rem}</style>
</head><body>

<nav class="navbar navbar-dark bg-dark mb-3">
 <div class="container-fluid">
   <a class="navbar-brand" href="/ka-part">🏢 아파트 관리</a>
   <div class="d-flex gap-2">
     <!-- ✅ 허브(상위 메뉴) -->
     <a class="btn btn-outline-light btn-sm" href="/ka-part">☰ 메뉴</a>

     <!-- ✅ 기존 4개(스샷과 동일) -->
     <a class="btn btn-outline-light btn-sm" href="{{ url_for('ui_home') }}">일검침</a>
     <a class="btn btn-outline-light btn-sm" href="{{ url_for('ui_files') }}">업무파일</a>
     <a class="btn btn-outline-light btn-sm" href="/ts/">tool-search</a>
     <a class="btn btn-outline-light btn-sm" href="{{ url_for('ui_complaints') }}">민원/고장</a>
     <a class="btn btn-outline-light btn-sm" href="{{ url_for('ui_settings') }}">설정</a>

     <!-- ✅ 공구(blueprint: /ts) -->
     <a class="btn btn-outline-light btn-sm" href="/ts">공구</a>
   </div>
 </div>
</nav>

<div class="container">
 {% with messages=get_flashed_messages() %}
   {% if messages %}<div class="alert alert-info">{{ messages[0] }}</div>{% endif %}
 {% endwith %}
 {{ body|safe }}
</div>
<script>
function fillFromSpeech(id){
  if(!('webkitSpeechRecognition' in window)){alert('이 브라우저는 음성입력이 지원되지 않습니다.');return;}
  const r = new webkitSpeechRecognition(); r.lang='ko-KR'; r.interimResults=false; r.maxAlternatives=1;
  r.onresult = e => { document.getElementById(id).value = e.results[0][0].transcript; };
  r.start();
}
</script>
<script>
if ("serviceWorker" in navigator) {
  window.addEventListener("load", function () {
    navigator.serviceWorker.register("/ka-part/sw.js").catch(function(){});
  });
}
</script>


</body></html>
"""
def render(title, body, **ctx):
    return render_template_string(BASE, title=title, body=body, **ctx)
UI_APPS = """
<div class="d-flex justify-content-between align-items-center mb-2">
  <h5 class="m-0">🏢 시설관리 메뉴</h5>
  <span class="text-muted small">현장=속도 · 서버=진실</span>
</div>

<div class="row g-2">
  {% for a in apps %}
  <div class="col-12 col-md-6">
    <a class="text-decoration-none" href="{{ a.href }}">
      <div class="card h-100">
        <div class="card-body">
          <div class="d-flex justify-content-between align-items-start">
            <div>
              <div class="fw-bold">{{ a.title }}</div>
              <div class="text-muted small mt-1">{{ a.desc }}</div>
            </div>
            <div class="text-muted">→</div>
          </div>
        </div>
      </div>
    </a>
  </div>
  {% endfor %}
</div>

<div class="alert alert-light border mt-3 small mb-0">
  기록(저장/접수/확정/삭제)은 서버가 정답입니다. 네트워크가 없으면 “보여주기”만 되고 “쓰기”는 실패로 남겨야 합니다.
</div>
"""

@app.route("/ka-part")
def ui_apps():
    apps = [
        {"title": "일검침/설비 일지", "desc": "전력/급수/열량/유량 기록·월별집계·CSV", "href": url_for("ui_home")},
        {"title": "업무 파일", "desc": "업로드·검색·카카오 전송(선택)", "href": url_for("ui_files")},
        {"title": "비교견적", "desc": "업체/품목/단가 매트릭스·DOCX 품의서", "href": url_for("ui_compare")},
        {"title": "민원/고장", "desc": "접수·분류·미디어 첨부", "href": url_for("ui_complaints")},
        {"title": "공구 검색/관리", "desc": "공구 모듈(blueprint: /ts)", "href": "/ts"},
        {"title": "설정", "desc": "보정계수/요금/카카오/공개URL", "href": url_for("ui_settings")},
    ]
    body = render_template_string(UI_APPS, apps=apps)
    return render("시설관리 메뉴", body)

# ───────────────────────────────────────────────────────────────────
# 8) UI: 일검침/설비 일지 목록 + CSV/월별/재계산
# ───────────────────────────────────────────────────────────────────
@app.route("/ui")
def ui_home():
    rows = SubstationLog.query.order_by(
        func.coalesce(SubstationLog.log_date, date(1900,1,1)).desc(),
        func.coalesce(SubstationLog.log_time, time(0,0)).desc(),
        SubstationLog.id.desc()
    ).limit(200).all()
    body = render_template_string("""
<div class="d-flex justify-content-between align-items-center mb-2">
  <h5 class="m-0">일검침/설비 일지</h5>
  <div>
    <a class="btn btn-sm btn-primary" href="{{ url_for('ui_new_log') }}">+ 새 기록</a>
    <a class="btn btn-sm btn-outline-secondary" href="{{ url_for('export_csv') }}">CSV</a>
    <a class="btn btn-sm btn-outline-secondary" href="{{ url_for('ui_monthly') }}">월별집계</a>
    <a class="btn btn-sm btn-outline-danger" href="{{ url_for('recalc_all') }}" onclick="return confirm('전체 재계산을 실행할까요?')">일괄 재계산</a>
  </div>
</div>
<div class="table-responsive">
<table class="table table-sm table-hover">
  <thead class="table-light">
  <tr>
    <th>ID</th><th>일시</th><th>vcb_Kw</th>
    <th>유효전력/사용량</th><th>산업용/사용량</th><th>가로등/사용량</th>
    <th>상수도/사용량</th><th>열량/사용량</th><th>유량/사용량</th>
    <th>비고</th><th></th>
  </tr>
  </thead>
  <tbody>
  {% for r in rows %}
    <tr>
      <td>{{ r.id }}</td>
      <td>{{ r.log_date }} {{ r.log_time or '' }}</td>
      <td>{{ r.electric_energy or '' }}</td>
      <td>{{ r.hv_acc_kwh or '' }}/<strong>{{ r.hv_use_kwh or '' }}</strong></td>
      <td>{{ r.ind_acc_kwh or '' }}/<strong>{{ r.ind_use_kwh or '' }}</strong></td>
      <td>{{ r.str_acc_kwh or '' }}/<strong>{{ r.str_use_kwh or '' }}</strong></td>
      <td>{{ r.acc_water or '' }}/<strong>{{ r.day_water or '' }}</strong></td>
      <td>{{ r.acc_heat or '' }}/<strong>{{ r.day_heat or '' }}</strong></td>
      <td>{{ r.acc_flow or '' }}/<strong>{{ r.day_flow or '' }}</strong></td>
      <td>{{ (r.event or '')[:12] }}</td>
      <td class="text-end">
        <a class="btn btn-sm btn-outline-primary" href="{{ url_for('ui_edit_log', lid=r.id) }}">수정</a>
        <form method="post" action="{{ url_for('del_log', lid=r.id) }}" style="display:inline" onsubmit="return confirm('삭제하시겠습니까?');">
          <button class="btn btn-sm btn-outline-danger">삭제</button>
        </form>
      </td>
    </tr>
  {% else %}
    <tr><td colspan="11" class="text-muted p-3">기록 없음</td></tr>
  {% endfor %}
  </tbody>
</table>
</div>
""", rows=rows)
    return render("일검침 일지", body)

# 신규/수정 폼 (폼 name이 모델 필드와 정확히 일치하도록 주의)
FORM = """
<div class="card"><div class="card-header"><strong>{{ title }}</strong></div>
<div class="card-body">
<form method="post">
<div class="row g-2">
  <!-- 날짜/시각/점검자 -->
  <div class="col-6 col-md-3"><label class="form-label">일자</label>
    <input type="date" name="log_date" class="form-control" value="{{v.log_date}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">시각</label>
    <input type="time" name="log_time" class="form-control" value="{{v.log_time}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">점검자</label>
   <select name="operator" class="form-select">
      {% for n in operators %}<option value="{{n}}" {% if v.operator==n %}selected{% endif %}>{{n}}</option>{% endfor %}
    </select>
  </div>
  
  <div class="col-3 col-md-2"><label class="form-label">vcb역률</label>
  <input name="vcb_p_factor" class="form-control" value="{{v.vcb_p_factor}}"></div>
  
  <div class="col-3 col-md-2"><label class="form-label">LV역률</label>
    <input name="power_factor" class="form-control" value="{{v.power_factor}}"></div>

  <!-- 고압 수전 측정치 (kV, A, PF) -->
  <div class="col-6 col-md-3"><label class="form-label">vcb(kV)</label>
    <input name="incomer_voltage" class="form-control" value="{{v.incomer_voltage}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">vcb(A)</label>
    <input name="incomer_curr" class="form-control" value="{{v.incomer_curr}}"></div>


  <!-- 저압 3회선 (V, A) + 공통 역률 -->
  <div class="col-6 col-md-3"><label class="form-label">LV1(V)</label>
    <input name="lv1_v" class="form-control" value="{{v.lv1_v}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">LV1(A)</label>
    <input name="lv1_a" class="form-control" value="{{v.lv1_a}}"></div>

  <div class="col-6 col-md-3"><label class="form-label">LV2(V)</label>
    <input name="lv2_v" class="form-control" value="{{v.lv2_v}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">LV2(A)</label>
    <input name="lv2_a" class="form-control" value="{{v.lv2_a}}"></div>

  <div class="col-6 col-md-3"><label class="form-label">LV3(V)</label>
    <input name="lv3_v" class="form-control" value="{{v.lv3_v}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">LV3(A)</label>
    <input name="lv3_a" class="form-control" value="{{v.lv3_a}}"></div>



  <!-- 전력 누적 계기값(당일 지시) -->
  <div class="col-4 col-md-2"><label class="form-label">유효전력</label>
    <input name="hv_acc_kwh" class="form-control" value="{{v.hv_acc_kwh}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">산 업</label>
    <input name="ind_acc_kwh" class="form-control" value="{{v.ind_acc_kwh}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">가로등</label>
    <input name="str_acc_kwh" class="form-control" value="{{v.str_acc_kwh}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">TR1</label><input name="TR1" class="form-control" value="{{v.TR1}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">TR2</label><input name="TR2" class="form-control" value="{{v.TR2}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">TR3</label><input name="TR3" class="form-control" value="{{v.TR3}}"></div>

  <!-- 설비 누적 값 -->
  <div class="col-4 col-md-2"><label class="form-label">상수도</label>
    <input name="acc_water" class="form-control" value="{{v.acc_water}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">열량</label>
    <input name="acc_heat" class="form-control" value="{{v.acc_heat}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">유량</label>
    <input name="acc_flow" class="form-control" value="{{v.acc_flow}}"></div>

  <!-- 온도 -->
  <div class="col-6 col-md-3"><label class="form-label">고층난방S</label><input name="hst" class="form-control" value="{{v.hst}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">고층난방R</label><input name="hrt" class="form-control" value="{{v.hrt}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">저층난방S</label><input name="lst" class="form-control" value="{{v.lst}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">저층난방R</label><input name="lrt" class="form-control" value="{{v.lrt}}"></div>

<!--<div class="col-6 col-md-3"><label class="form-label">권선온도</label><input name="winding_temp" class="form-control" value="{{v.winding_temp}}"></div> -->

  <div class="col-12"><label class="form-label">특이사항</label>
    <input id="event" name="event" class="form-control" value="{{v.event}}">
    <button type="button" class="btn btn-sm btn-outline-secondary mt-1" onclick="fillFromSpeech('event')">🎤 음성으로 입력</button>
  </div>
  <div class="col-12"><label class="form-label">비고</label>
    <textarea name="remarks" class="form-control" rows="2">{{v.remarks}}</textarea></div>
</div>
<div class="mt-3 d-flex gap-2">
  <button class="btn btn-primary">저장</button>
  <a class="btn btn-secondary" href="{{ url_for('ui_home') }}">목록</a>
</div>
</form>
</div></div>
"""

@app.route("/ui/new", methods=["GET","POST"])
def ui_new_log():
    """새 기록 생성(모든 칼럼 저장) + 자동계산"""
    if request.method == "POST":
        s = Settings.get()
        payload = {k: request.form.get(k) for k in request.form.keys()}

        # ① 폼 → 모델 매핑 (모든 칼럼 저장)
        r = SubstationLog(
            log_date=parse_date(payload.get("log_date")) or date.today(),
            log_time=parse_time(payload.get("log_time")),
            operator=payload.get("operator") or DEFAULT_OPERATOR,

            incomer_voltage=parse_float(payload.get("incomer_voltage")),
            incomer_curr=parse_float(payload.get("incomer_curr")),
            vcb_p_factor=parse_float(payload.get("vcb_p_factor")),

            lv1_v=parse_float(payload.get("lv1_v")), lv1_a=parse_float(payload.get("lv1_a")),
            lv2_v=parse_float(payload.get("lv2_v")), lv2_a=parse_float(payload.get("lv2_a")),
            lv3_v=parse_float(payload.get("lv3_v")), lv3_a=parse_float(payload.get("lv3_a")),
            power_factor=parse_float(payload.get("power_factor")),

            hv_acc_kwh=parse_float(payload.get("hv_acc_kwh")),
            ind_acc_kwh=parse_float(payload.get("ind_acc_kwh")),
            str_acc_kwh=parse_float(payload.get("str_acc_kwh")),

            acc_water=parse_float(payload.get("acc_water")),
            acc_heat=parse_float(payload.get("acc_heat")),
            acc_flow=parse_float(payload.get("acc_flow")),

            hst=parse_float(payload.get("hst")), hrt=parse_float(payload.get("hrt")),
            lst=parse_float(payload.get("lst")), lrt=parse_float(payload.get("lrt")),
            TR1=parse_float(payload.get("TR1")), TR2=parse_float(payload.get("TR2")),

            TR3=parse_float(payload.get("TR3")),
            winding_temp=parse_float(payload.get("winding_temp")),
            event=payload.get("event") or "", remarks=payload.get("remarks") or ""
        )

        # ② 전일 기록 조회(일 사용량 계산용) — 같은 달/연속 불문, 단순 직전 날짜
        prev = SubstationLog.query\
            .filter(SubstationLog.log_date < r.log_date)\
            .order_by(SubstationLog.log_date.desc(), SubstationLog.id.desc())\
            .first()

        # ③ 자동 계산(수전/저압 kW, 일 사용량)
        compute_auto_fields(r, prev, s)

        db.session.add(r); db.session.commit()
        flash(f"등록 완료 (ID {r.id})")
        return redirect(url_for("ui_home"))

    # GET: 폼 초기값
    v = dict(
        log_date=date.today().isoformat(), log_time=datetime.now().strftime("%H:%M"),
        operator=DEFAULT_OPERATOR, incomer_voltage="", incomer_curr="", vcb_p_factor="",
        lv1_v="", lv1_a="", lv2_v="", lv2_a="", lv3_v="", lv3_a="", power_factor="",
        hv_acc_kwh="", ind_acc_kwh="", str_acc_kwh="",
        acc_water="", acc_heat="", acc_flow="",
        hst="", hrt="", lst="", lrt="", TR1="", TR2="", TR3="", winding_temp="",
        event="", remarks=""
    )
    return render("새 기록", render_template_string(FORM, title="새 기록", v=v, operators=OPERATORS))

@app.route("/ui/edit/<int:lid>", methods=["GET","POST"])
def ui_edit_log(lid):
    """기존 기록 수정(모든 칼럼 저장) + 자동 재계산"""
    r = SubstationLog.query.get_or_404(lid)
    if request.method == "POST":
        s = Settings.get()

        # ① 폼 → 필드 갱신 (누락 없이 전부)
        r.log_date = parse_date(request.form.get("log_date")) or r.log_date
        r.log_time = parse_time(request.form.get("log_time"))
        r.operator  = request.form.get("operator") or DEFAULT_OPERATOR

        r.incomer_voltage = parse_float(request.form.get("incomer_voltage"))
        r.incomer_curr    = parse_float(request.form.get("incomer_curr"))
        r.vcb_p_factor    = parse_float(request.form.get("vcb_p_factor"))

        r.lv1_v = parse_float(request.form.get("lv1_v")); r.lv1_a = parse_float(request.form.get("lv1_a"))
        r.lv2_v = parse_float(request.form.get("lv2_v")); r.lv2_a = parse_float(request.form.get("lv2_a"))
        r.lv3_v = parse_float(request.form.get("lv3_v")); r.lv3_a = parse_float(request.form.get("lv3_a"))
        r.power_factor = parse_float(request.form.get("power_factor"))

        r.hv_acc_kwh  = parse_float(request.form.get("hv_acc_kwh"))
        r.ind_acc_kwh = parse_float(request.form.get("ind_acc_kwh"))
        r.str_acc_kwh = parse_float(request.form.get("str_acc_kwh"))

        r.acc_water = parse_float(request.form.get("acc_water"))
        r.acc_heat  = parse_float(request.form.get("acc_heat"))
        r.acc_flow  = parse_float(request.form.get("acc_flow"))

        r.hst = parse_float(request.form.get("hst")); r.hrt = parse_float(request.form.get("hrt"))
        r.lst = parse_float(request.form.get("lst")); r.lrt = parse_float(request.form.get("lrt"))
        r.TR1 = parse_float(request.form.get("TR1")); r.TR2 = parse_float(request.form.get("TR2"))

        r.TR3 = parse_float(request.form.get("TR3"))
        r.winding_temp = parse_float(request.form.get("winding_temp"))
        r.event = request.form.get("event",""); r.remarks = request.form.get("remarks","")

        # ② 수정된 날짜 기준, 이전 레코드로 일사용량 재계산
        prev = SubstationLog.query\
            .filter(SubstationLog.log_date < r.log_date)\
            .order_by(SubstationLog.log_date.desc(), SubstationLog.id.desc())\
            .first()

        # ③ 자동 계산 반영
        compute_auto_fields(r, prev, s)

        db.session.commit()
        flash("수정 완료")
        return redirect(url_for("ui_home"))

    # GET: 폼에 현재 값 표시
    v = r.to_dict()
    return render("기록 수정", render_template_string(FORM, title=f"기록 수정 #{lid}", v=v, operators=OPERATORS))

@app.route("/ui/del/<int:lid>", methods=["POST"])
def del_log(lid):
    """기록 삭제"""
    r = SubstationLog.query.get_or_404(lid)
    db.session.delete(r); db.session.commit()
    flash("삭제되었습니다.")
    return redirect(url_for("ui_home"))

# CSV 내보내기
@app.route("/export.csv")
def export_csv():
    """모든 기록 CSV 저장 후 다운로드"""
    path = os.path.join(app.instance_path, f"substation_{date.today().isoformat()}.csv")
    import csv
    rows = SubstationLog.query.order_by(SubstationLog.log_date.asc(), SubstationLog.id.asc()).all()
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["id","date","time","operator",
                    "elec_kw","hv_acc","hv_use","ind_acc","ind_use","str_acc","str_use",
                    "water_acc","water_day","heat_acc","heat_day","flow_acc","flow_day",
                    "event","remarks"])
        for r in rows:
            w.writerow([
                r.id, r.log_date, r.log_time, r.operator,
                r3(r.electric_energy), r3(r.hv_acc_kwh), r3(r.hv_use_kwh),
                r3(r.ind_acc_kwh), r3(r.ind_use_kwh), r3(r.str_acc_kwh), r3(r.str_use_kwh),
                r3(r.acc_water), r3(r.day_water), r3(r.acc_heat), r3(r.day_heat),
                r3(r.acc_flow), r3(r.day_flow), r.event, r.remarks
            ])
    return send_file(path, as_attachment=True, download_name=os.path.basename(path))

# 월별 집계
@app.route("/monthly")
def ui_monthly():
    """연/월로 일사용량 합계를 보여줌"""
    y = int(request.args.get("y", date.today().year))
    m = int(request.args.get("m", date.today().month))
    start = date(y, m, 1)
    end = date(y + (m==12), (m%12)+1, 1)
    rows = SubstationLog.query.filter(SubstationLog.log_date >= start, SubstationLog.log_date < end).all()
    sums = dict(
        hv=sum([r.hv_use_kwh or 0 for r in rows]),
        ind=sum([r.ind_use_kwh or 0 for r in rows]),
        st=sum([r.str_use_kwh or 0 for r in rows]),
        water=sum([r.day_water or 0 for r in rows]),
        heat=sum([r.day_heat or 0 for r in rows]),
        flow=sum([r.day_flow or 0 for r in rows]),
    )
    body = render_template_string("""
<h5>{{y}}-{{'%02d'%m}} 월별 집계</h5>
<table class="table table-sm w-auto">
<tr><th>고압 사용량</th><td>{{ '%.3f' % sums.hv }}</td></tr>
<tr><th>산업용 사용량</th><td>{{ '%.3f' % sums.ind }}</td></tr>
<tr><th>가로등 사용량</th><td>{{ '%.3f' % sums.st }}</td></tr>
<tr><th>급수 일사용 합</th><td>{{ '%.3f' % sums.water }}</td></tr>
<tr><th>열량 일사용 합</th><td>{{ '%.3f' % sums.heat }}</td></tr>
<tr><th>유량 일사용 합</th><td>{{ '%.3f' % sums.flow }}</td></tr>
</table>
<p><a class="btn btn-secondary" href="{{ url_for('ui_home') }}">← 돌아가기</a></p>
""", y=y, m=m, sums=type("Obj",(object,),sums))
    return render("월별 집계", body)

# 일괄 재계산
@app.route("/recalc")
def recalc_all():
    """전체 레코드에 대해 자동 계산(일사용량 포함)을 재적용"""
    s = Settings.get()
    rows = SubstationLog.query.order_by(SubstationLog.log_date.asc(), SubstationLog.id.asc()).all()
    prev = None
    for r in rows:
        compute_auto_fields(r, prev, s)
        prev = r
    db.session.commit()
    flash("전체 재계산 완료")
    return redirect(url_for("ui_home"))

# ───────────────────────────────────────────────────────────────────
# 9) UI: 업무 파일 보관/검색/삭제
# ───────────────────────────────────────────────────────────────────
# 9-추가) UI: 비교견적 (목록/입력/집계/워드 출력)
# ───────────────────────────────────────────────────────────────────
def _fmt_won(v):
    try:
        n = float(v or 0)
        return f"{int(round(n)):,.0f}원"
    except:
        return ""

def _ensure_price(vendor_id, item_id):
    row = VendorPrice.query.filter_by(vendor_id=vendor_id, item_id=item_id).first()
    if not row:
        row = VendorPrice(vendor_id=vendor_id, item_id=item_id, unit_price=None)
        db.session.add(row); db.session.commit()
    return row

@app.route("/compare")
def ui_compare():
    ensure_compare_schema()
    sets = CompareSet.query.order_by(CompareSet.created_at.desc(), CompareSet.id.desc()).all()
    body = render_template_string(SUBTABS + """
<div class="d-flex justify-content-between align-items-center mb-2">
  <h5 class="m-0">⚡ 비교견적</h5>
  <a class="btn btn-sm btn-primary" href="{{ url_for('ui_compare_new') }}">+ 새 비교</a>
</div>
<div class="list-group">
{% for s in sets %}
  <a class="list-group-item list-group-item-action" href="{{ url_for('ui_compare_edit', sid=s.id) }}">
    <div class="d-flex justify-content-between">
      <div>
        <strong>{{ s.title }}</strong> <span class="text-muted small">#{{ s.id }} · {{ s.created_at.strftime("%Y-%m-%d") }}</span>
        <div class="small text-muted">{{ (s.memo or '')[:120] }}</div>
      </div>
      <div class="text-end">
        <span class="badge bg-secondary">VAT {{ ((s.vat_rate or 0)*100) | int }}%</span>
      </div>
    </div>
  </a>
{% else %}
  <div class="text-muted p-3">비교 세트가 없습니다.</div>
{% endfor %}
</div>
""", SUBTABS=SUBTABS)
    return render("비교견적", body)

@app.route("/compare/new", methods=["GET","POST"])
def ui_compare_new():
    if request.method=="POST":
        title = (request.form.get("title") or "").strip() or "무제 비교"
        requester = request.form.get("requester") or ""
        memo = request.form.get("memo") or ""
        vat = parse_float(request.form.get("vat_rate")) or 0.1
        s = CompareSet(title=title, requester=requester, memo=memo, vat_rate=vat)
        db.session.add(s); db.session.commit()
        flash("생성되었습니다.")
        return redirect(url_for("ui_compare_edit", sid=s.id))
    body = render_template_string(SUBTABS + """
<div class="card"><div class="card-header"><strong>새 비교</strong></div>
<div class="card-body">
  <form method="post" class="row g-2">
    <div class="col-12 col-md-6"><label class="form-label">제목</label><input class="form-control" name="title" required></div>
    <div class="col-6 col-md-3"><label class="form-label">의뢰자</label><input class="form-control" name="requester"></div>
    <div class="col-6 col-md-3"><label class="form-label">VAT율(예: 0.1)</label><input class="form-control" name="vat_rate" value="0.1"></div>
    <div class="col-12"><label class="form-label">메모</label><textarea class="form-control" name="memo" rows="2"></textarea></div>
    <div class="col-12 mt-2"><button class="btn btn-primary">생성</button> <a class="btn btn-secondary" href="{{ url_for('ui_compare') }}">목록</a></div>
  </form>
</div></div>
""", SUBTABS=SUBTABS)
    return render("새 비교", body)

@app.route("/compare/<int:sid>", methods=["GET","POST"])
def ui_compare_edit(sid):
    s = CompareSet.query.get_or_404(sid)
    # 상단 설정 수정
    if request.method=="POST" and request.form.get("_form")=="base":
        s.title = request.form.get("title") or s.title
        s.requester = request.form.get("requester") or s.requester
        s.memo = request.form.get("memo") or s.memo
        s.vat_rate = parse_float(request.form.get("vat_rate")) or s.vat_rate
        db.session.commit(); flash("저장되었습니다.")
        return redirect(url_for("ui_compare_edit", sid=sid))

    vendors = Vendor.query.filter_by(compare_id=sid).order_by(Vendor.id.asc()).all()
    items = Item.query.filter_by(compare_id=sid).order_by(Item.id.asc()).all()

    # 가격 매트릭스 구성
    price = {}
    for v in vendors:
        price[v.id] = {}
    for vp in VendorPrice.query.join(Vendor, Vendor.id==VendorPrice.vendor_id)\
                               .filter(Vendor.compare_id==sid).all():
        price.setdefault(vp.vendor_id, {})[vp.item_id] = vp.unit_price

    # 총액 집계
    sum_supply = {v.id: 0.0 for v in vendors}
    for it in items:
        for v in vendors:
            up = price.get(v.id, {}).get(it.id)
            if up is not None and it.qty:
                sum_supply[v.id] += float(up) * float(it.qty)
    sum_vat = {vid: (sum_supply[vid]*(s.vat_rate or 0.0)) for vid in sum_supply}
    sum_total = {vid: (sum_supply[vid]+sum_vat[vid]) for vid in sum_supply}

    # 최저가(총액) 벤더
    best_vendor_id = None
    if vendors:
        avail = [(vid, total) for vid,total in sum_total.items() if total>0]
        if avail:
            best_vendor_id = min(avail, key=lambda x: x[1])[0]

    body = render_template_string(SUBTABS + """
<div class="d-flex justify-content-between align-items-center mb-2">
  <h5 class="m-0">⚡ 비교편집 — #{{s.id}}</h5>
  <div class="d-flex gap-2">
    <a class="btn btn-sm btn-outline-secondary" href="{{ url_for('ui_compare') }}">목록</a>
    <a class="btn btn-sm btn-success" href="{{ url_for('compare_export_docx', sid=s.id) }}">품의서(DOCX)</a>
  </div>
</div>

<div class="card mb-3"><div class="card-header"><strong>기본정보</strong></div>
<div class="card-body">
<form method="post" class="row g-2">
  <input type="hidden" name="_form" value="base">
  <div class="col-12 col-md-6"><label class="form-label">제목</label>
    <input class="form-control" name="title" value="{{s.title}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">의뢰자</label>
    <input class="form-control" name="requester" value="{{s.requester or ''}}"></div>
  <div class="col-6 col-md-3"><label class="form-label">VAT율</label>
    <input class="form-control" name="vat_rate" value="{{s.vat_rate}}"></div>
  <div class="col-12"><label class="form-label">메모</label>
    <textarea class="form-control" name="memo" rows="2">{{s.memo or ''}}</textarea></div>
  <div class="col-12 mt-2"><button class="btn btn-primary">저장</button></div>
</form>
</div></div>

<div class="row g-3">
  <div class="col-12 col-lg-4">
    <div class="card h-100"><div class="card-header"><strong>업체</strong></div>
      <div class="card-body">
        <form class="row g-2 mb-2" method="post" action="{{ url_for('compare_add_vendor', sid=s.id) }}">
          <div class="col-12"><input class="form-control" name="name" placeholder="업체명" required></div>
          <div class="col-6"><input class="form-control" name="contact" placeholder="담당자"></div>
          <div class="col-6"><input class="form-control" name="phone" placeholder="연락처"></div>
          <div class="col-12"><input class="form-control" name="note" placeholder="비고"></div>
          <div class="col-12"><button class="btn btn-outline-primary w-100">+ 추가</button></div>
        </form>
        <div class="list-group">
        {% for v in vendors %}
          <div class="list-group-item d-flex justify-content-between {% if v.id==best_vendor_id %}list-group-item-success{% endif %}">
            <div>
              <strong>{{ v.name }}</strong> <span class="small text-muted">{{ v.contact }} {{ v.phone }}</span>
              {% if v.id==best_vendor_id %}<span class="badge bg-success">최저가</span>{% endif %}
            </div>
            <form method="post" action="{{ url_for('compare_del_vendor', sid=s.id, vid=v.id) }}" onsubmit="return confirm('삭제?');">
              <button class="btn btn-sm btn-outline-danger">삭제</button>
            </form>
          </div>
        {% else %}
          <div class="text-muted small p-2">업체 없음</div>
        {% endfor %}
        </div>
      </div>
    </div>
  </div>

  <div class="col-12 col-lg-8">
    <div class="card"><div class="card-header"><strong>품목 & 단가</strong></div>
      <div class="card-body">
        <form class="row g-2 mb-3" method="post" action="{{ url_for('compare_add_item', sid=s.id) }}">
          <div class="col-5 col-md-4"><input class="form-control" name="name" placeholder="품명" required></div>
          <div class="col-7 col-md-4"><input class="form-control" name="spec" placeholder="규격/설명"></div>
          <div class="col-4 col-md-2"><input class="form-control" name="unit" value="EA" placeholder="단위"></div>
          <div class="col-4 col-md-2"><input class="form-control" name="qty" type="number" step="0.01" value="1"></div>
          <div class="col-12"><button class="btn btn-outline-primary w-100">+ 품목추가</button></div>
        </form>

        <div class="table-responsive">
        <form method="post" action="{{ url_for('compare_update_prices', sid=s.id) }}">
          <table class="table table-sm table-bordered align-middle">
            <thead class="table-light">
              <tr>
                <th style="min-width:120px">품명</th><th>규격</th>
                <th class="text-end">수량</th><th>단위</th>
                {% for v in vendors %}<th class="text-end">{{ v.name }}<div class="small text-muted">단가(공급가)</div></th>{% endfor %}
              </tr>
            </thead>
            <tbody>
            {% for it in items %}
              <tr>
                <td>{{ it.name }}</td>
                <td class="small">{{ it.spec or '' }}</td>
                <td class="text-end">{{ '%.2f' % it.qty }}</td>
                <td>{{ it.unit }}</td>
                {% for v in vendors %}
                  {% set val = price.get(v.id, {}).get(it.id) %}
                  <td><input class="form-control form-control-sm text-end" name="p_{{ v.id }}_{{ it.id }}" value="{{ '' if val is none else ('%.0f' % val) }}" placeholder="-"></td>
                {% endfor %}
              </tr>
            {% else %}
              <tr><td colspan="{{ 4 + (vendors|length) }}" class="text-muted">품목 없음</td></tr>
            {% endfor %}
            </tbody>
          </table>
          <div class="d-flex justify-content-between">
            <div>
              <a class="btn btn-sm btn-outline-danger" href="{{ url_for('compare_clear_prices', sid=s.id) }}" onclick="return confirm('모든 단가를 비우시겠습니까?');">단가 초기화</a>
            </div>
            <div><button class="btn btn-primary">단가 저장</button></div>
          </div>
        </form>
        </div>

        <hr>
        <h6>집계</h6>
        <div class="table-responsive">
          <table class="table table-sm w-auto">
            <thead class="table-light"><tr><th>업체</th><th class="text-end">공급가 합계</th><th class="text-end">부가세</th><th class="text-end">총액</th></tr></thead>
            <tbody>
              {% for v in vendors %}
              <tr class="{% if v.id==best_vendor_id %}table-success{% endif %}">
                <td><strong>{{ v.name }}</strong>{% if v.id==best_vendor_id %} <span class="badge bg-success">최저가</span>{% endif %}</td>
                <td class="text-end">{{ _fmt_won(sum_supply[v.id]) }}</td>
                <td class="text-end">{{ _fmt_won(sum_vat[v.id]) }}</td>
                <td class="text-end"><strong>{{ _fmt_won(sum_total[v.id]) }}</strong></td>
              </tr>
              {% endfor %}
            </tbody>
          </table>
        </div>

      </div>
    </div>
  </div>
</div>
""", SUBTABS=SUBTABS,
    s=s, vendors=vendors, items=items, price=price,
    sum_supply=sum_supply, sum_vat=sum_vat, sum_total=sum_total,
    best_vendor_id=best_vendor_id, _fmt_won=_fmt_won)
    return render("비교편집", body)

@app.route("/compare/<int:sid>/vendor/add", methods=["POST"])
def compare_add_vendor(sid):
    CompareSet.query.get_or_404(sid)
    v = Vendor(compare_id=sid,
               name=(request.form.get("name") or "업체").strip(),
               contact=request.form.get("contact") or "",
               phone=request.form.get("phone") or "",
               note=request.form.get("note") or "")
    db.session.add(v); db.session.commit()
    flash("업체 추가")
    return redirect(url_for("ui_compare_edit", sid=sid))

@app.route("/compare/<int:sid>/vendor/<int:vid>/del", methods=["POST"])
def compare_del_vendor(sid, vid):
    v = Vendor.query.get_or_404(vid)
    if v.compare_id != sid: flash("잘못된 요청"); return redirect(url_for("ui_compare_edit", sid=sid))
    # 가격행 먼저 정리
    VendorPrice.query.filter_by(vendor_id=vid).delete()
    db.session.delete(v); db.session.commit()
    flash("업체 삭제")
    return redirect(url_for("ui_compare_edit", sid=sid))

@app.route("/compare/<int:sid>/item/add", methods=["POST"])
def compare_add_item(sid):
    CompareSet.query.get_or_404(sid)
    qty = parse_float(request.form.get("qty")) or 1.0
    it = Item(compare_id=sid,
              name=(request.form.get("name") or "품목").strip(),
              spec=request.form.get("spec") or "",
              unit=(request.form.get("unit") or "EA").strip(),
              qty=qty)
    db.session.add(it); db.session.commit()
    # 기존 업체들에 대해 가격행 보장
    for v in Vendor.query.filter_by(compare_id=sid).all():
        _ensure_price(v.id, it.id)
    flash("품목 추가")
    return redirect(url_for("ui_compare_edit", sid=sid))

@app.route("/compare/<int:sid>/prices/save", methods=["POST"])
def compare_update_prices(sid):
    vendors = Vendor.query.filter_by(compare_id=sid).all()
    items = Item.query.filter_by(compare_id=sid).all()
    # form 키: p_{vendor_id}_{item_id}
    for v in vendors:
        for it in items:
            key = f"p_{v.id}_{it.id}"
            val = request.form.get(key)
            if val is not None:
                vp = _ensure_price(v.id, it.id)
                vp.unit_price = parse_float(val)
    db.session.commit()
    flash("단가 저장 완료")
    return redirect(url_for("ui_compare_edit", sid=sid))

@app.route("/compare/<int:sid>/prices/clear")
def compare_clear_prices(sid):
    VendorPrice.query.join(Vendor, Vendor.id==VendorPrice.vendor_id)\
                     .filter(Vendor.compare_id==sid).update({VendorPrice.unit_price: None})
    db.session.commit()
    flash("모든 단가를 비웠습니다.")
    return redirect(url_for("ui_compare_edit", sid=sid))

# DOCX 품의서 출력 (python-docx 없으면 CSV로 대체)
@app.route("/compare/<int:sid>/export.docx")
def compare_export_docx(sid):
    s = CompareSet.query.get_or_404(sid)
    vendors = Vendor.query.filter_by(compare_id=sid).order_by(Vendor.id.asc()).all()
    items = Item.query.filter_by(compare_id=sid).order_by(Item.id.asc()).all()
    price = {}
    for vp in VendorPrice.query.join(Vendor, Vendor.id==VendorPrice.vendor_id)\
                               .filter(Vendor.compare_id==sid).all():
        price.setdefault(vp.vendor_id, {})[vp.item_id] = vp.unit_price

    def sums():
        sum_supply = {v.id: 0.0 for v in vendors}
        for it in items:
            for v in vendors:
                up = price.get(v.id, {}).get(it.id)
                if up is not None and it.qty:
                    sum_supply[v.id] += float(up)*float(it.qty)
        sum_vat = {vid: (sum_supply[vid]*(s.vat_rate or 0.0)) for vid in sum_supply}
        sum_total = {vid: (sum_supply[vid]+sum_vat[vid]) for vid in sum_supply}
        return sum_supply, sum_vat, sum_total
    sum_supply, sum_vat, sum_total = sums()
    best = None
    if vendors:
        avail = [(v.id, sum_total[v.id]) for v in vendors if sum_total[v.id]>0]
        if avail: best = min(avail, key=lambda x:x[1])[0]

    # 파일 경로
    out_name = f"compare_{sid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(UPLOAD_DIR, out_name)

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        doc = Document()
        doc.add_heading(f'품의서 - {s.title}', level=1)
        p = doc.add_paragraph()
        p.add_run(f"의뢰자: {s.requester or '-'}    VAT: {int((s.vat_rate or 0.0)*100)}%\n").font.size = Pt(10)
        if s.memo:
            doc.add_paragraph(s.memo).italic = True

        # 표: 헤더(품명/규격/수량/단위 + 업체별 단가)
        cols = 4 + len(vendors)
        table = doc.add_table(rows=1, cols=cols)
        hdr = table.rows[0].cells
        hdr[0].text = "품명"; hdr[1].text="규격"; hdr[2].text="수량"; hdr[3].text="단위"
        for i,v in enumerate(vendors, start=4):
            hdr[i].text = f"{v.name}\n(단가)"
        for it in items:
            row = table.add_row().cells
            row[0].text = it.name
            row[1].text = it.spec or ""
            row[2].text = f"{it.qty:g}"
            row[3].text = it.unit
            for i,v in enumerate(vendors, start=4):
                up = price.get(v.id, {}).get(it.id)
                row[i].text = (f"{int(round(up)):,}" if up is not None else "-")

        doc.add_paragraph("")
        # 집계
        table2 = doc.add_table(rows=1+len(vendors), cols=4)
        table2.rows[0].cells[0].text = "업체"
        table2.rows[0].cells[1].text = "공급가 합계"
        table2.rows[0].cells[2].text = "부가세"
        table2.rows[0].cells[3].text = "총액"
        for i,v in enumerate(vendors, start=1):
            table2.rows[i].cells[0].text = v.name + (" (최저가)" if best==v.id else "")
            table2.rows[i].cells[1].text = _fmt_won(sum_supply[v.id])
            table2.rows[i].cells[2].text = _fmt_won(sum_vat[v.id])
            table2.rows[i].cells[3].text = _fmt_won(sum_total[v.id])

        doc.add_paragraph("\n결론: 상기 비교 결과를 참조하여 최적 업체를 선정하고자 합니다.")
        doc.save(out_path)
        # WorkFile에 등재
        st = os.stat(out_path)
        wf = WorkFile(title=f"품의서-{s.title}", description=f"비교견적 DOCX (세트 #{s.id})",
                      category="품의서", tags="비교견적,품의",
                      filename=os.path.basename(out_path), ext=".docx", size=st.st_size, uploader="관리자")
        db.session.add(wf); db.session.commit()
        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))
    except Exception as e:
        # 폴백: CSV 내보내기
        csv_name = f"compare_{sid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        csv_path = os.path.join(UPLOAD_DIR, csv_name)
        import csv
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            hdr = ["품명","규격","수량","단위"] + [f"{v.name}(단가)" for v in vendors]
            w.writerow(hdr)
            for it in items:
                row = [it.name, it.spec or "", it.qty, it.unit]
                for v in vendors:
                    up = price.get(v.id, {}).get(it.id)
                    row.append("" if up is None else int(round(up)))
                w.writerow(row)
            w.writerow([])
            w.writerow(["업체","공급가 합","부가세","총액"])
            for v in vendors:
                w.writerow([v.name, int(round(sum_supply[v.id])), int(round(sum_vat[v.id])), int(round(sum_total[v.id]))])
        st = os.stat(csv_path)
        wf = WorkFile(title=f"품의서CSV-{s.title}", description=f"비교견적 CSV (DOCX 실패 폴백) 세트 #{s.id}",
                      category="품의서", tags="비교견적,품의",
                      filename=os.path.basename(csv_path), ext=".csv", size=st.st_size, uploader="관리자")
        db.session.add(wf); db.session.commit()
        flash("python-docx 사용 불가로 CSV로 내보냈습니다.")
        return send_file(csv_path, as_attachment=True, download_name=os.path.basename(csv_path))
@app.route("/files", methods=["GET","POST"])
def ui_files():
    """업무 파일 업로드/검색/삭제"""
    if request.method == "POST":
        f = request.files.get("file")
        title = request.form.get("title") or (f.filename if f else "무제")
        cat = request.form.get("category") or "일반"
        tags = request.form.get("tags") or ""
        desc = request.form.get("description") or ""
        if not f:
            flash("파일이 없습니다."); return redirect(url_for("ui_files"))
        ext = os.path.splitext(f.filename)[1].lower()
        new_name = f"{uuid.uuid4().hex}{ext}"
        path = os.path.join(UPLOAD_DIR, new_name)
        f.save(path)
        st = os.stat(path)
        wf = WorkFile(title=title, description=desc, category=cat, tags=tags,
                      filename=new_name, ext=ext, size=st.st_size, uploader="관리자")
        db.session.add(wf); db.session.commit()
        flash("업로드 완료")
        return redirect(url_for("ui_files"))

    q = request.args.get("q","").strip()
    query = WorkFile.query
    if q:
        like = f"%{q}%"
        query = query.filter(
            (WorkFile.title.like(like)) | (WorkFile.description.like(like)) |
            (WorkFile.category.like(like)) | (WorkFile.tags.like(like))
        )
    rows = query.order_by(WorkFile.created_at.desc(), WorkFile.id.desc()).all()
    body = render_template_string("""
    <h5 class="mb-2">업무 파일</h5>
<form class="row g-2 mb-3" method="post" enctype="multipart/form-data">

<div class="mb-3">
  <ul class="nav nav-pills">
    <li class="nav-item">
      <a class="nav-link {% if request.path.startswith('/files') %}active{% endif %}"
         href="{{ url_for('ui_files') }}">📂 파일보관</a>
    </li>
    <li class="nav-item">
      <a class="nav-link {% if request.path.startswith('/compare') %}active{% endif %}"
         href="{{ url_for('ui_compare') }}">⚡ 비교견적</a>
    </li>
  </ul>
</div>


  <div class="col-12 col-md-3"><input class="form-control" name="title" placeholder="제목"></div>
  <div class="col-6 col-md-2"><input class="form-control" name="category" placeholder="분류"></div>
  <div class="col-6 col-md-3"><input class="form-control" name="tags" placeholder="태그(,구분)"></div>
  <div class="col-12"><input class="form-control" name="description" placeholder="설명"></div>
  <div class="col-8"><input class="form-control" type="file" name="file" required></div>
  <div class="col-4"><button class="btn btn-primary w-100">업로드</button></div>
</form>

<form class="input-group mb-2" method="get">
  <input class="form-control" name="q" value="{{ request.args.get('q','') }}" placeholder="검색(제목/설명/분류/태그)">
  <button class="btn btn-outline-secondary">검색</button>
</form>

<form id="kakao-files-form" method="post" action="{{ url_for('kakao_send_files') }}"></form>

<form class="d-flex gap-2 mb-2" method="post" action="{{ url_for('kakao_send_files') }}">
  <input class="form-control" name="message" placeholder="전송 메모(선택)">
  <button class="btn btn-warning">선택 항목 카카오 전송</button>
</form>

<div class="list-group">
{% for r in rows %}
  <div class="list-group-item">
    <div class="d-flex justify-content-between">
      <div>
        <!-- ✅ 체크박스: kakao-files-form 으로 제출 -->
        <input class="form-check-input me-2" type="checkbox" name="fid" form="kakao-files-form" value="{{ r.id }}">
        <strong>{{ r.title }}</strong>
        <span class="text-muted">[{{ r.category }}] {{ r.tags }}</span>
        <div class="small">{{ r.description }}</div>
      </div>
      <div class="text-end">
        <a class="btn btn-sm btn-outline-secondary" href="{{ url_for('download_file', fid=r.id) }}">다운로드</a>
        <form method="post" action="{{ url_for('delete_file', fid=r.id) }}" style="display:inline" onsubmit="return confirm('삭제?');">
          <button class="btn btn-sm btn-outline-danger">삭제</button>
        </form>
      </div>
    </div>
  </div>
{% else %}
  <div class="text-muted p-3">파일 없음</div>
{% endfor %}
</div>
""",rows=rows)
    return render("업무 파일", body)

@app.route("/kakao/send/files", methods=["POST"])
def kakao_send_files():
    s = Settings.get()
    if not s.kakao_access_token or not s.kakao_friend_uuid:
        flash("설정에서 Kakao Access Token과 Friend UUID를 입력하세요.")
        return redirect(url_for("ui_files"))

    ids = request.form.getlist("fid")
    if not ids:
        # 체크박스가 숨은 폼으로 들어오는 케이스 대비: 파일 목록 폼에서 message만 온 경우
        ids = request.form.getlist("fid[]")
    ids = [int(x) for x in ids if str(x).isdigit()]
    if not ids:
        flash("전송할 파일을 선택하세요.")
        return redirect(url_for("ui_files"))

    rows = WorkFile.query.filter(WorkFile.id.in_(ids)).order_by(WorkFile.id.asc()).all()
    msg = (request.form.get("message") or "").strip()

    # 메시지 본문 구성
    lines = []
    if msg: lines.append(f"[메모] {msg}")
    lines.append("📎 업무파일 전송")
    pub = s.public_base_url or ""
    for r in rows:
        link = (pub + r.filename) if pub else None
        base = f"- {r.title}{r.ext or ''}"
        lines.append(base + (f" 🔗 {link}" if link else ""))

    text = "\n".join(lines)[:990]
    uuids = [x.strip() for x in (s.kakao_friend_uuid or "").split(",") if x.strip()]
    ok, detail = kakao_send_default(s.kakao_access_token, uuids, text, link_url=s.public_base_url or None)
    flash("카카오 전송 성공" if ok else f"카카오 전송 실패: {detail}")
    return redirect(url_for("ui_files"))

@app.route("/files/<int:fid>/download")
def download_file(fid):
    r = WorkFile.query.get_or_404(fid)
    path = os.path.join(UPLOAD_DIR, r.filename)
    return send_file(path, as_attachment=True, download_name=r.title + (r.ext or ""))

@app.route("/files/<int:fid>/delete", methods=["POST"])
def delete_file(fid):
    r = WorkFile.query.get_or_404(fid)
    try:
        os.remove(os.path.join(UPLOAD_DIR, r.filename))
    except Exception:
        pass
    db.session.delete(r); db.session.commit()
    flash("삭제되었습니다.")
    return redirect(url_for("ui_files"))

# ───────────────────────────────────────────────────────────────────
# 10) UI: 민원/고장 (간단 자동분류 + 미디어 첨부)
# ───────────────────────────────────────────────────────────────────
COMPLAINT_CATEGORIES = {
    "전기": ["정전", "누전", "조명", "차단기", "콘센트", "승강기"],
    "배관": ["누수", "막힘", "수압", "배수", "악취"],
    "난방": ["보일러", "온수", "난방", "온도"],
    "시설": ["문고장", "파손", "도색", "청소"],
}
DEFAULT_CATEGORY = "기타"
ALL_TAGS = set(sum(COMPLAINT_CATEGORIES.values(), []))

def simple_classify(text: str, filename: Optional[str]) -> (str, float, List[str]):
    """아주 단순한 규칙 기반 분류(키워드 카운트) + 미디어 첨부 힌트"""
    txt = (text or "").lower()
    score = {}; tags = []
    for cat, keys in COMPLAINT_CATEGORIES.items():
        s = sum([1 for k in keys if k.lower() in txt])
        score[cat] = s
        if s: tags.extend([k for k in keys if k.lower() in txt])
    if filename:
        ext = os.path.splitext(filename)[1].lower()
        if ext in (".jpg",".jpeg",".png",".mp4",".mov",".avi"):
            tags.append("미디어첨부")
    best = max(score, key=score.get) if score else DEFAULT_CATEGORY
    conf = (score.get(best,0) / max(1,len(ALL_TAGS))) + (0.1 if "미디어첨부" in tags else 0.0)
    conf = r3(min(conf, 1.0))
    return best if score.get(best,0)>0 else DEFAULT_CATEGORY, conf, tags

@app.route("/c", methods=["GET","POST"])
def ui_complaints():
    """민원/고장 접수 + 목록"""
    if request.method == "POST":
        name = request.form.get("name") or "익명"
        unit = request.form.get("unit") or ""
        phone = request.form.get("phone") or ""
        textv = request.form.get("text") or ""
        ch = request.form.get("channel") or "웹"

        media = request.files.get("media")
        media_name, media_type = None, None
        if media and media.filename:
            ext = os.path.splitext(media.filename)[1].lower()
            media_name = f"cmp_{uuid.uuid4().hex}{ext}"
            media.save(os.path.join(UPLOAD_DIR, media_name))
            mt = mimetypes.guess_type(media_name)[0] or ""
            media_type = "image" if mt.startswith("image") else ("video" if mt.startswith("video") else "file")

        cat, conf, tags = simple_classify(textv, media_name)
        row = Complaint(name=name, unit=unit, phone=phone, text=textv, channel=ch,
                        category=cat, confidence=conf, tags=",".join(tags),
                        media_filename=media_name, media_type=media_type)
        db.session.add(row); db.session.commit()
        flash(f"접수 완료 (#{row.id}, {row.category})")
        return redirect(url_for("ui_complaints"))

    q = request.args.get("q","").strip()
    query = Complaint.query
    if q:
        like = f"%{q}%"
        query = query.filter( (Complaint.text.like(like)) | (Complaint.category.like(like)) | (Complaint.unit.like(like)) )
    rows = query.order_by(Complaint.created_at.desc(), Complaint.id.desc()).all()
    body = render_template_string("""
<h5 class="mb-2">민원/고장 접수</h5>
<form class="row g-2 mb-3" method="post" enctype="multipart/form-data">
  <div class="col-6 col-md-2"><input class="form-control" name="name" placeholder="이름"></div>
  <div class="col-6 col-md-2"><input class="form-control" name="unit" placeholder="동/호수"></div>
  <div class="col-6 col-md-2"><input class="form-control" name="phone" placeholder="연락처"></div>
  <div class="col-6 col-md-2"><input class="form-control" name="channel" value="웹"></div>
  <div class="col-12"><textarea id="ctext" class="form-control" name="text" rows="2" placeholder="내용"></textarea>
    <button type="button" class="btn btn-sm btn-outline-secondary mt-1" onclick="fillFromSpeech('ctext')">🎤 음성으로 입력</button>
  </div>
  <div class="col-8"><input class="form-control" type="file" name="media" accept="image/*,video/*"></div>
  <div class="col-4"><button class="btn btn-primary w-100">접수</button></div>
</form>

<form class="input-group mb-2" method="get">
  <input class="form-control" name="q" value="{{ request.args.get('q','') }}" placeholder="검색(내용/분류/호수)">
  <button class="btn btn-outline-secondary">검색</button>
</form>

<div class="list-group">
{% for r in rows %}
  <div class="list-group-item">
    <div class="d-flex justify-content-between">
      <div><strong>#{{ r.id }}</strong> [{{ r.category }}] <span class="text-muted small">{{ r.unit }} {{ r.name }}</span>
        <div class="small text-muted">{{ r.created_at.strftime("%Y-%m-%d %H:%M") }}</div>
        <div>{{ r.text }}</div>
        {% if r.media_filename %}
          <div class="small">첨부: <a href="{{ url_for('download_upload', name=r.media_filename) }}">{{ r.media_filename }}</a></div>
        {% endif %}
      </div>
      <div class="text-end">
        <span class="badge bg-secondary">{{ r.status }}</span>
      </div>
    </div>
  </div>
{% else %}
  <div class="text-muted p-3">접수 없음</div>
{% endfor %}
</div>
""", rows=rows)
    return render("민원/고장", body)

@app.route("/u/<path:name>")
def download_upload(name):
    """업로드 파일 내려받기(민원 첨부 포함)"""
    path = os.path.join(UPLOAD_DIR, name)
    return send_file(path, as_attachment=True)

# # ───────────────────────────────────────────────────────────────────
# Kakao 메시지 전송 헬퍼
# ───────────────────────────────────────────────────────────────────
def kakao_send_default(access_token: str, friend_uuids: list[str], text: str, link_url: str | None = None):
    """
    카카오 친구에게 기본 템플릿(텍스트) 메시지 전송.
    - friend_uuids: ["uuid1","uuid2", ...]
    - text: 본문 (최대 200자 권장)
    - link_url: 버튼 링크(선택). public_base_url이 있을 때 파일/첨부 링크로 사용 가능
    요구:
      - 카카오 개발자 콘솔 애플리케이션
      - 액세스토큰에 friends, talk_message 권한
      - 수신자 uuid 확보 (friends API 등)
    """
    url = "https://kapi.kakao.com/v1/api/talk/friends/message/default/send"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/x-www-form-urlencoded;charset=utf-8",
    }
    template_obj = {
        "object_type": "text",
        "text": text[:990],  # 안전 여유
        "link": {"web_url": link_url or "https://developers.kakao.com"},
        "button_title": "열기" if link_url else "확인",
    }
    data = {
        "receiver_uuids": json.dumps(friend_uuids, ensure_ascii=False),
        "template_object": json.dumps(template_obj, ensure_ascii=False),
    }
    resp = requests.post(url, headers=headers, data=data, timeout=10)
    ok = (resp.status_code == 200)
    return ok, (resp.text if not ok else "OK")
# 11) UI: 설정 (보정계수/요금/카카오)
# ───────────────────────────────────────────────────────────────────
@app.route("/ka-part/manifest.webmanifest")
def manifest():
    # static/manifest.webmanifest 파일을 그대로 서빙
    path = os.path.join(app.static_folder, "manifest.webmanifest")
    return send_file(path, mimetype="application/manifest+json", max_age=0)

@app.route("/ka-part/sw.js")
def sw():
    path = os.path.join(app.static_folder, "sw.js")
    return send_file(path, mimetype="application/javascript", max_age=0)

@app.route("/ka-part/ui")
def ui_home_alias():
    return ui_home()

@app.route("/ka-part/files", methods=["GET","POST"])
def ui_files_alias():
    return ui_files()

@app.route("/ka-part/compare")
def ui_compare_alias():
    return ui_compare()

@app.route("/ka-part/c", methods=["GET","POST"])
def ui_complaints_alias():
    return ui_complaints()

@app.route("/ka-part/settings", methods=["GET","POST"])
def ui_settings_alias():
    return ui_settings()


@app.route("/settings", methods=["GET","POST"])
def ui_settings():
    """보정계수/요금/카카오 전송 설정"""
    s = Settings.get()
    if request.method == "POST":
        # 빈칸은 기존값 유지 (parse_float(None) 방지)
        def keep(val, cur):
            p = parse_float(val)
            return p if p is not None else cur

        s.hv_factor  = keep(request.form.get("hv_factor"), s.hv_factor)
        s.ind_factor = keep(request.form.get("ind_factor"), s.ind_factor)
        s.street_factor = keep(request.form.get("street_factor"), s.street_factor)
        s.water_factor = keep(request.form.get("water_factor"), s.water_factor)
        s.heat_factor  = keep(request.form.get("heat_factor"), s.heat_factor)
        s.flow_factor  = keep(request.form.get("flow_factor"), s.flow_factor)

        s.tariff_per_kwh = keep(request.form.get("tariff_per_kwh"), s.tariff_per_kwh)
        s.base_charge    = keep(request.form.get("base_charge"), s.base_charge)
        s.allocation_method = request.form.get("allocation_method") or s.allocation_method

        s.kakao_rest_key     = request.form.get("kakao_rest_key") or s.kakao_rest_key
        s.kakao_access_token = request.form.get("kakao_access_token") or s.kakao_access_token
        s.kakao_friend_uuid  = request.form.get("kakao_friend_uuid") or s.kakao_friend_uuid

        # ✅ 이 줄은 반드시 커밋/리디렉트 전에 있어야 적용됩니다.
        s.public_base_url = request.form.get("public_base_url") or s.public_base_url

        db.session.commit()
        flash("저장되었습니다.")
        return redirect(url_for("ui_settings"))

    body = render_template_string("""
<h5>설정</h5>
<form method="post" class="row g-2">
  <div class="col-4 col-md-2"><label class="form-label">고압계수</label><input class="form-control" name="hv_factor" value="{{s.hv_factor}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">산업계수</label><input class="form-control" name="ind_factor" value="{{s.ind_factor}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">가로등계수</label><input class="form-control" name="street_factor" value="{{s.street_factor}}"></div>

  <div class="col-4 col-md-2"><label class="form-label">급수계수</label><input class="form-control" name="water_factor" value="{{s.water_factor}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">열량계수</label><input class="form-control" name="heat_factor" value="{{s.heat_factor}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">유량계수</label><input class="form-control" name="flow_factor" value="{{s.flow_factor}}"></div>

  <div class="col-4 col-md-2"><label class="form-label">kWh 단가</label><input class="form-control" name="tariff_per_kwh" value="{{s.tariff_per_kwh}}"></div>
  <div class="col-4 col-md-2"><label class="form-label">기본요금</label><input class="form-control" name="base_charge" value="{{s.base_charge}}"></div>
  <div class="col-4 col-md-3"><label class="form-label">배분방식</label><input class="form-control" name="allocation_method" value="{{s.allocation_method}}"></div>

  <div class="col-12"><hr></div>
  <div class="col-12"><strong>카카오 전송(선택)</strong></div>
  <div class="col-12 col-md-4"><label class="form-label">REST Key</label><input class="form-control" name="kakao_rest_key" value="{{s.kakao_rest_key or ''}}"></div>
  <div class="col-12 col-md-4"><label class="form-label">Access Token</label><input class="form-control" name="kakao_access_token" value="{{s.kakao_access_token or ''}}"></div>
  <div class="col-12 col-md-4"><label class="form-label">Friend UUID</label><input class="form-control" name="kakao_friend_uuid" value="{{s.kakao_friend_uuid or ''}}"></div>

  <div class="col-12"><hr></div>
  <div class="col-12 col-md-6">
    <label class="form-label">공개 URL Prefix(선택)</label>
    <input class="form-control" name="public_base_url" value="{{s.public_base_url or ''}}" placeholder="예: https://files.example.com/apt/">
    <div class="form-text">여기에 설정하면 파일/첨부 전송 시 해당 URL로 링크를 붙여 보냅니다.</div>
  </div>

  <div class="col-12 mt-2">
    <button class="btn btn-primary">저장</button>
    <a class="btn btn-secondary" href="{{ url_for('ui_home') }}">돌아가기</a>
  </div>
</form>
""", s=s)
    return render("설정", body)

# ───────────────────────────────────────────────────────────────────
# 12) 서버 시작: 테이블 생성→컬럼 보강→설정 1행 보장→실행
# ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        # ✅ SQLite면 누락 컬럼 보강, Postgres면 create_all로 충분
        auto_migrate_columns()
        Settings.get()
    app.run(host="127.0.0.1", port=8000, debug=False, use_reloader=False, threaded=False)
